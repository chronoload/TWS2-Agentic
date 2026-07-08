#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工具系统 - 基于 CoreCoder 设计
包含文件操作、编辑、搜索等常用工具

优化内容:
- 结构化结果支持 (ToolResult dataclass)
- 结果大小限制与自动截断
- JSON/human-readable 双模式输出
"""

import os
import re
import sys
import json
import time
import logging
import locale
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Callable
from abc import ABC, abstractmethod
from dataclasses import dataclass, field, asdict

logger = logging.getLogger(__name__)

MAX_TOOL_RESULT_CHARS = 999999


def _get_system_encoding() -> str:
    """获取系统首选编码"""
    try:
        return locale.getpreferredencoding(False) or "utf-8"
    except Exception:
        return "utf-8"


def _read_file_with_encoding(file_path: Path, fallback_encodings: List[str] = None) -> str:
    """智能读取文件，自动检测编码"""
    if fallback_encodings is None:
        fallback_encodings = ["utf-8", "gbk", "gb2312", "gb18030"]
    
    encodings_to_try = [_get_system_encoding()]
    for enc in fallback_encodings:
        if enc not in encodings_to_try:
            encodings_to_try.append(enc)
    
    for encoding in encodings_to_try:
        try:
            return file_path.read_text(encoding=encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    
    try:
        return file_path.read_text(errors="replace")
    except Exception:
        return ""


def safe_truncate(text: str, max_chars: int, suffix: str = None) -> tuple:
    """
    安全截断文本，避免在格式中间截断
    
    返回: (截断后文本, 是否截断, 原始长度)
    """
    if not text or len(text) <= max_chars:
        return text, False, len(text)
    
    suffix = suffix or f"\n\n... [已截断, 原始 {len(text)} 字符]"
    suffix_len = len(suffix)
    available = max_chars - suffix_len - 50
    
    if available <= 0:
        return text[:max_chars] + "...", True, len(text)
    
    truncated_content = text[:available]
    
    lines = truncated_content.split('\n')
    if len(lines) > 1:
        lines = lines[:-1]
        truncated_content = '\n'.join(lines)
        if len(truncated_content) > max_chars - len(suffix) - 20:
            truncated_content = truncated_content[:max_chars - len(suffix) - 20]
    
    result = truncated_content.strip() + suffix
    return result, True, len(text)


@dataclass
class ToolResult:
    """结构化工具结果 - 可解析, 支持 JSON 和文本双模式"""
    success: bool = True
    data: Dict[str, Any] = field(default_factory=dict)
    message: str = ""
    error: str = ""
    truncated: bool = False
    original_length: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)

    def __str__(self) -> str:
        """人类可读的字符串表示"""
        parts = []
        if self.message:
            parts.append(self.message)
        if self.error:
            parts.append(f"❌ {self.error}")
        if self.truncated:
            parts.append(f"\n⚠️ 结果过长，已截断（原始 {self.original_length} 字符）")
        return "\n".join(parts)

    def to_json(self) -> str:
        """JSON 表示 - 可被下游解析"""
        return json.dumps(asdict(self), ensure_ascii=False)

    @classmethod
    def ok(cls, data: Optional[Dict[str, Any]] = None, message: str = "",
           metadata: Optional[Dict[str, Any]] = None) -> "ToolResult":
        return cls(success=True, data=data or {}, message=message,
                   metadata=metadata or {})

    @classmethod
    def err(cls, error: str, data: Optional[Dict[str, Any]] = None) -> "ToolResult":
        return cls(success=False, error=error, data=data or {})

    @classmethod
    def from_text(cls, text: str, max_chars: int = MAX_TOOL_RESULT_CHARS) -> "ToolResult":
        """从纯文本创建, 自动安全截断"""
        original_length = len(text)
        if original_length <= max_chars:
            message = text
            truncated = False
        else:
            message, truncated, _ = safe_truncate(text, max_chars)
        return cls(success=True, message=message, truncated=truncated,
                   original_length=original_length,
                   data={"text": message, "original_length": original_length})


class Tool(ABC):
    name: str
    description: str
    parameters: Dict[str, Any]
    max_result_chars: int = MAX_TOOL_RESULT_CHARS
    category: str = "general"
    keywords: List[str] = []
    risk_level: str = "medium"
    requires_context: List[str] = []
    model_hint: str = ""  # 增强格式: "[何时使用] 简短描述\n[参数说明]\n- param: 用法"

    @abstractmethod
    def execute(self, **kwargs) -> str:
        pass

    def execute_structured(self, **kwargs) -> ToolResult:
        text_result = self.execute(**kwargs)
        return ToolResult.from_text(text_result, self.max_result_chars)

    def schema(self) -> Dict[str, Any]:
        schema = {
            "type": "function",
            "function": {
                "name": self.name,
                "description": self.description,
                "parameters": self.parameters,
            },
        }
        # 注入 model_hint 到 description 末尾（model_hint 本身已带前缀）
        if self.model_hint:
            schema["function"]["description"] = f"{self.description}\n\n{self.model_hint}"
        return schema

    def search_signature(self) -> str:
        return f"{self.name} {self.description} {' '.join(self.keywords)} {self.category}"

    def _truncate_error(self, text: str, max_chars: Optional[int] = None) -> str:
        """截断错误信息"""
        limit = max_chars or self.max_result_chars
        if len(text) <= limit:
            return text
        return text[:limit] + f"\n\n... [已截断, 原始 {len(text)} 字符]"


class ReadFileTool(Tool):
    name = "read_file"
    category = "file_io"
    keywords = ["read", "file", "文件", "读取", "查看", "代码"]
    model_hint = "[何时使用] 查看文件内容、分析代码。\n[参数说明]\n- path: 必填，文件路径\n- lines: 只看前N行（适合大文件头部预览）\n- tail: 只看末尾N行（适合日志）\n- start_line/end_line: 指定行范围\n- max_chars: 截断长度控制"
    description = "读取文件内容，支持全量/前N行/范围/尾部读取。"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "【必填】文件路径（相对于工作目录）", "example": "src/main.py"},
            "encoding": {"type": "string", "description": "文件编码（默认 utf-8）"},
            "lines": {"type": "integer", "description": "读取前 N 行内容（适合快速查看大文件头部）"},
            "start_line": {"type": "integer", "description": "读取起始行号（从1开始），需与 end_line 同时使用", "example": 10},
            "end_line": {"type": "integer", "description": "读取结束行号（包含），需与 start_line 同时使用", "example": 50},
            "tail": {"type": "integer", "description": "读取末尾 N 行内容（适合查看日志文件最新追加的行）"},
            "max_chars": {"type": "integer", "description": "最大返回字符数，超出则截断（默认 8000）"},
        },
        "required": ["path"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, path: str, encoding: str = "utf-8", 
                lines: Optional[int] = None,
                start_line: Optional[int] = None,
                end_line: Optional[int] = None,
                tail: Optional[int] = None,
                max_chars: Optional[int] = None) -> str:
        try:
            full_path = self.base_dir / Path(path)
            if not full_path.exists():
                return ToolResult.err(f"文件不存在：{path}").to_json()
            
            content = full_path.read_text(encoding=encoding)
            all_lines = content.splitlines()
            total_lines = len(all_lines)
            
            read_lines = all_lines
            
            # 处理不同读取模式
            mode_desc = ""
            if tail is not None and tail > 0:
                read_lines = all_lines[-tail:]
                mode_desc = f"尾部{tail}行"
            elif start_line is not None and end_line is not None:
                start = max(0, start_line - 1)
                end = min(total_lines, end_line)
                read_lines = all_lines[start:end]
                mode_desc = f"第{start_line}-{end_line}行"
            elif lines is not None and lines > 0:
                read_lines = all_lines[:lines]
                mode_desc = f"前{lines}行"
            else:
                mode_desc = f"全部{total_lines}行"
            
            display_content = "\n".join(read_lines)
            
            # 处理字符截断（安全截断）
            char_limit = max_chars or self.max_result_chars
            original_len = len(display_content)
            if original_len > char_limit:
                display_content, truncated, _ = safe_truncate(display_content, char_limit)
            else:
                truncated = False
            
            result = f"--- {path} ({mode_desc}) ---\n{display_content}"
            return ToolResult(
                success=True,
                data={
                    "path": path,
                    "total_lines": total_lines,
                    "read_lines": len(read_lines),
                    "read_range": mode_desc,
                    "content": display_content,
                    "original_length": original_len,
                    "truncated": truncated,
                },
                message=result,
                truncated=truncated,
            ).to_json()
        except Exception as e:
            return ToolResult.err(f"读取文件失败：{e}").to_json()

    def execute_structured(self, path: str, encoding: str = "utf-8",
                         lines: Optional[int] = None,
                         start_line: Optional[int] = None,
                         end_line: Optional[int] = None,
                         tail: Optional[int] = None,
                         max_chars: Optional[int] = None) -> ToolResult:
        try:
            full_path = self.base_dir / Path(path)
            if not full_path.exists():
                return ToolResult.err(f"文件不存在：{path}")
            
            content = full_path.read_text(encoding=encoding)
            all_lines = content.splitlines()
            total_lines = len(all_lines)
            
            read_lines = all_lines
            mode_desc = ""
            
            if tail is not None and tail > 0:
                read_lines = all_lines[-tail:]
                mode_desc = f"尾部{tail}行"
            elif start_line is not None and end_line is not None:
                start = max(0, start_line - 1)
                end = min(total_lines, end_line)
                read_lines = all_lines[start:end]
                mode_desc = f"第{start_line}-{end_line}行"
            elif lines is not None and lines > 0:
                read_lines = all_lines[:lines]
                mode_desc = f"前{lines}行"
            else:
                mode_desc = f"全部{total_lines}行"
            
            display_content = "\n".join(read_lines)
            original_content_length = len(display_content)
            char_limit = max_chars or self.max_result_chars
            truncated = original_content_length > char_limit
            if truncated:
                display_content = display_content[:char_limit] + f"\n\n... [已截断, 原始 {original_content_length} 字符]"
            
            return ToolResult(
                success=True,
                data={
                    "path": path,
                    "total_lines": total_lines,
                    "read_lines": len(read_lines),
                    "read_range": mode_desc,
                    "content": display_content,
                    "original_length": original_content_length,
                    "truncated": truncated,
                },
                message=f"--- {path} ({mode_desc}) ---\n{display_content}",
                truncated=truncated,
                original_length=original_content_length,
            )
        except Exception as e:
            return ToolResult.err(f"读取文件失败：{e}")


class WriteFileTool(Tool):
    name = "write_file"
    category = "file_io"
    keywords = ["write", "file", "文件", "写入", "保存", "创建"]
    max_result_chars = 99999999
    model_hint = "[何时使用] 创建新文件或完全覆盖内容。小改动请用 edit_file。\n[参数说明]\n- path: 必填，文件路径\n- content: 必填，完整内容\n- mode: create=新建覆盖，append=追加到末尾\n- open_after: 写入后自动用编辑器打开"
    description = "创建或覆盖文件。小修改建议用 edit_file，大改动用 write_file。"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "【必填】文件路径（相对于工作目录）", "example": "src/new_module.py"},
            "content": {"type": "string", "description": "【必填】要写入的完整文件内容。此参数为写入的整个文件内容，不可省略。如果你只想追加几行内容，请使用 edit_file 工具的 append 模式。"},
            "encoding": {"type": "string", "description": "文件编码（默认 utf-8）"},
            "mode": {"type": "string", "description": "写入模式：create=新建或完全覆盖，append=追加到文件末尾", "enum": ["create", "append"], "default": "create", "example": "create"},
            "open_after": {"type": "boolean", "description": "写入后用编辑器打开预览（默认 false）", "default": False},
        },
        "required": ["path", "content"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, path: str, content: str, encoding: str = "utf-8", mode: str = "create", open_after: bool = False) -> str:
        result = self.execute_structured(path, content, encoding, mode, open_after)
        return result.to_json()

    def execute_structured(self, path: str, content: str, encoding: str = "utf-8", mode: str = "create", open_after: bool = False) -> ToolResult:
        try:
            full_path = self.base_dir / Path(path)
            full_path.parent.mkdir(parents=True, exist_ok=True)

            if mode == "append":
                existing = ""
                if full_path.exists():
                    existing = full_path.read_text(encoding=encoding)
                    if existing and not existing.endswith('\n'):
                        existing += '\n'
                full_path.write_text(existing + content, encoding=encoding)
                total_lines = (existing + content).splitlines()
                appended_lines = content.splitlines()
                message = f"✅ 已追加到 {path} ({len(appended_lines)} 行追加，总计 {len(total_lines)} 行)"
            else:
                full_path.write_text(content, encoding=encoding)
                line_count = len(content.splitlines())
                message = f"✅ 已写入 {path} ({line_count} 行)"

            if open_after:
                try:
                    import subprocess
                    import platform
                    system = platform.system()
                    if system == "Windows":
                        os.startfile(str(full_path))
                    elif system == "Darwin":
                        subprocess.run(["open", str(full_path)], check=True)
                    else:
                        subprocess.run(["xdg-open", str(full_path)], check=True)
                    message += "\n📖 已在默认编辑器中打开文件"
                except Exception as e:
                    message += f"\n⚠️ 尝试打开文件时出错: {e}"

            return ToolResult.ok(
                data={"path": path, "mode": mode, "open_after": open_after},
                message=message
            )
        except Exception as e:
            return ToolResult.err(f"写入文件失败：{e}")


class EditFileTool(Tool):
    name = "edit_file"
    category = "file_io"
    keywords = ["edit", "file", "修改", "编辑", "替换"]
    model_hint = "[何时使用] 小范围修改文件内容。大改动/新建请用 write_file。\n[参数说明]\n- path: 必填，文件路径\n- mode: replace(搜索替换)/insert(行前插入)/append(末尾追加)/delete_lines(删除行范围)/undo(撤销上次编辑)\n- old_str: [replace] 被替换的旧文本\n- new_str: [replace/insert/append] 新文本\n- line_number: [insert] 在此行前插入\n- start_line/end_line: [delete_lines] 删除行范围\n- backup: 自动创建 .bak 备份"
    description = "编辑文件：替换、插入、追加、删除行或撤销上次编辑。"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "【必填】文件路径", "example": "src/main.py"},
            "mode": {
                "type": "string",
                "enum": ["replace", "insert", "append", "delete_lines", "undo"],
                "description": "编辑模式（默认 replace）：replace=搜索替换（需 old_str+new_str）, insert=指定行前插入（需 line_number+new_str）, append=末尾追加（需 new_str）, delete_lines=删除行范围（需 start_line+end_line）, undo=撤销上次编辑",
                "default": "replace",
                "example": "replace"
            },
            "old_str": {"type": "string", "description": "【replace 模式必填】要被替换的旧文本，必须精确匹配文件中的连续行"},
            "new_str": {"type": "string", "description": "【replace/insert/append 模式必填】替换或追加的新内容"},
            "line_number": {"type": "integer", "description": "【insert 模式必填】在此行号前插入新内容（从1开始）", "example": 10},
            "start_line": {"type": "integer", "description": "【delete_lines 模式必填】删除起始行号（从1开始）", "example": 5},
            "end_line": {"type": "integer", "description": "【delete_lines 模式必填】删除结束行号", "example": 15},
            "encoding": {"type": "string", "description": "文件编码（默认 utf-8）", "default": "utf-8"},
            "backup": {"type": "boolean", "description": "是否创建 .bak 备份文件（默认 false）", "default": False},
            "open_after": {"type": "boolean", "description": "编辑后用编辑器打开预览（默认 false）", "default": False},
        },
        "required": ["path"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, **kwargs) -> str:
        result = self.execute_structured(**kwargs)
        return result.to_json()

    def execute_structured(self, path: str, mode: str = "replace",
                           old_str: Optional[str] = None, new_str: Optional[str] = None,
                           line_number: Optional[int] = None,
                           start_line: Optional[int] = None, end_line: Optional[int] = None,
                           encoding: str = "utf-8", backup: bool = False,
                           open_after: bool = False) -> ToolResult:
        try:
            full_path = self.base_dir / Path(path)

            if mode == "undo":
                result = self._do_undo(full_path, path, encoding)
            else:
                if not full_path.exists():
                    return ToolResult.err(f"文件不存在：{path}")

                content = full_path.read_text(encoding=encoding)
                all_lines = content.splitlines(keepends=True)
                total_lines = len(all_lines)

                if mode == "replace":
                    result = self._do_replace(full_path, path, content, all_lines, total_lines,
                                              old_str, new_str, encoding, backup)
                elif mode == "insert":
                    result = self._do_insert(full_path, path, content, all_lines, total_lines,
                                           new_str, line_number, encoding, backup)
                elif mode == "append":
                    result = self._do_append(full_path, path, content, total_lines,
                                           new_str, encoding, backup)
                elif mode == "delete_lines":
                    result = self._do_delete_lines(full_path, path, content, all_lines, total_lines,
                                                 start_line, end_line, encoding, backup)
                else:
                    return ToolResult.err(f"未知编辑模式：{mode}，支持 replace/insert/append/delete_lines/undo")
            
            if open_after and result.success:
                try:
                    import subprocess
                    import platform
                    system = platform.system()
                    if system == "Windows":
                        os.startfile(str(full_path))
                    elif system == "Darwin":
                        subprocess.run(["open", str(full_path)], check=True)
                    else:
                        subprocess.run(["xdg-open", str(full_path)], check=True)
                    result.message += "\n📖 已在默认编辑器中打开文件"
                except Exception as e:
                    result.message += f"\n⚠️ 尝试打开文件时出错: {e}"
            
            return result
        except Exception as e:
            return ToolResult.err(f"编辑文件失败：{e}")

    def _do_replace(self, full_path, path, content, all_lines, total_lines,
                     old_str, new_str, encoding, backup) -> ToolResult:
        if not old_str:
            return ToolResult.err("replace 模式需要 old_str 参数")
        if new_str is None:
            new_str = ""

        count = content.count(old_str)
        if count == 0:
            return ToolResult.err("未找到匹配内容")
        if count > 1:
            return ToolResult.err(f"找到 {count} 处匹配，请更精确地选择替换内容")

        match_pos = content.index(old_str)
        match_line = content[:match_pos].count('\n') + 1

        if backup:
            self._create_backup(full_path)

        new_content = content.replace(old_str, new_str)
        full_path.write_text(new_content, encoding=encoding)

        diff = self._generate_unified_diff(old_str, new_str, match_line)
        context = self._get_context(all_lines, total_lines, match_line, new_str.count('\n') + 1)

        return ToolResult.ok(
            data={
                "path": path,
                "mode": "replace",
                "replacements": 1,
                "match_line": match_line,
                "diff": diff,
                "context": context,
            },
            message=f"✅ 已替换 {path} (第{match_line}行)\n--- 变更 ---\n{diff}\n--- 上下文 ---\n{context}"
        )

    def _do_insert(self, full_path, path, content, all_lines, total_lines,
                    new_str, line_number, encoding, backup) -> ToolResult:
        if new_str is None:
            return ToolResult.err("insert 模式需要 new_str 参数")
        if line_number is None:
            return ToolResult.err("insert 模式需要 line_number 参数")
        if line_number < 1 or line_number > total_lines + 1:
            return ToolResult.err(f"行号超出范围：{line_number}，文件共 {total_lines} 行")

        if backup:
            self._create_backup(full_path)

        insert_text = new_str if new_str.endswith('\n') else new_str + '\n'
        insert_idx = line_number - 1
        all_lines.insert(insert_idx, insert_text)
        new_content = ''.join(all_lines)
        full_path.write_text(new_content, encoding=encoding)

        diff = self._generate_unified_diff("", new_str, line_number)
        context = self._get_context(all_lines, len(all_lines), line_number, new_str.count('\n') + 1)

        return ToolResult.ok(
            data={
                "path": path,
                "mode": "insert",
                "inserted_at_line": line_number,
                "lines_inserted": new_str.count('\n') + 1,
                "total_lines": len(all_lines),
                "diff": diff,
                "context": context,
            },
            message=f"✅ 已在第{line_number}行前插入内容到 {path} ({new_str.count(chr(10))+1}行)\n--- 变更 ---\n{diff}\n--- 上下文 ---\n{context}"
        )

    def _do_append(self, full_path, path, content, total_lines,
                    new_str, encoding, backup) -> ToolResult:
        if new_str is None:
            return ToolResult.err("append 模式需要 new_str 参数")

        if backup:
            self._create_backup(full_path)

        append_text = new_str
        if content and not content.endswith('\n'):
            append_text = '\n' + append_text
        if not append_text.endswith('\n'):
            append_text += '\n'

        new_content = content + append_text
        full_path.write_text(new_content, encoding=encoding)

        new_total = new_content.count('\n') + (1 if not new_content.endswith('\n') else 0)
        start_line = total_lines + 1

        diff = self._generate_unified_diff("", new_str, start_line)

        return ToolResult.ok(
            data={
                "path": path,
                "mode": "append",
                "appended_at_line": start_line,
                "lines_appended": new_str.count('\n') + 1,
                "total_lines": new_total,
                "diff": diff,
            },
            message=f"✅ 已追加内容到 {path} (从第{start_line}行开始, {new_str.count(chr(10))+1}行)\n--- 变更 ---\n{diff}"
        )

    def _do_delete_lines(self, full_path, path, content, all_lines, total_lines,
                          start_line, end_line, encoding, backup) -> ToolResult:
        if start_line is None or end_line is None:
            return ToolResult.err("delete_lines 模式需要 start_line 和 end_line 参数")
        if start_line < 1 or end_line > total_lines or start_line > end_line:
            return ToolResult.err(f"行范围无效：{start_line}-{end_line}，文件共 {total_lines} 行")

        if backup:
            self._create_backup(full_path)

        deleted_text = ''.join(all_lines[start_line - 1:end_line])
        del all_lines[start_line - 1:end_line]
        new_content = ''.join(all_lines)
        full_path.write_text(new_content, encoding=encoding)

        diff = self._generate_unified_diff(deleted_text, "", start_line)
        context = self._get_context(all_lines, len(all_lines), max(1, start_line - 1), 3)

        return ToolResult.ok(
            data={
                "path": path,
                "mode": "delete_lines",
                "deleted_range": f"{start_line}-{end_line}",
                "lines_deleted": end_line - start_line + 1,
                "total_lines": len(all_lines),
                "diff": diff,
                "context": context,
            },
            message=f"✅ 已删除 {path} 第{start_line}-{end_line}行 ({end_line-start_line+1}行)\n--- 变更 ---\n{diff}\n--- 上下文 ---\n{context}"
        )

    def _create_backup(self, full_path: Path):
        backup_path = full_path.with_suffix(full_path.suffix + '.bak')
        import shutil
        shutil.copy2(full_path, backup_path)

    def _do_undo(self, full_path, path, encoding) -> ToolResult:
        backup_path = full_path.with_suffix(full_path.suffix + '.bak')
        if not backup_path.exists():
            return ToolResult.err(f"没有找到备份文件：{backup_path.name}")
        try:
            import shutil
            shutil.copy2(str(backup_path), str(full_path))
            content = full_path.read_text(encoding=encoding)
            line_count = len(content.splitlines())
            return ToolResult.ok(
                data={"path": path, "mode": "undo", "restored_lines": line_count},
                message=f"✅ 已从备份恢复 {path} ({line_count} 行)"
            )
        except Exception as e:
            return ToolResult.err(f"恢复备份失败：{e}")

    def _get_context(self, all_lines, total_lines, center_line, context_radius=3):
        start = max(0, center_line - context_radius - 1)
        end = min(total_lines, center_line + context_radius)
        lines = all_lines[start:end]
        result = []
        for i, line in enumerate(lines):
            line_no = start + i + 1
            marker = ">>>" if line_no == center_line else "   "
            result.append(f"{marker} {line_no:4d} | {line.rstrip()}")
        return '\n'.join(result)

    def _generate_unified_diff(self, old: str, new: str, start_line: int = 1) -> str:
        old_lines = old.splitlines(keepends=True)
        new_lines = new.splitlines(keepends=True)
        result = []
        for i, line in enumerate(old_lines):
            result.append(f"- {start_line + i:4d} | {line.rstrip()}")
        if old_lines and new_lines:
            result.append("      ---")
        for i, line in enumerate(new_lines):
            result.append(f"+ {start_line + i:4d} | {line.rstrip()}")
        return '\n'.join(result)


class ListDirectoryTool(Tool):
    name = "list_directory"
    category = "file_io"
    keywords = ["list", "directory", "目录", "文件夹", "ls"]
    model_hint = "查看目录结构或文件列表时使用。"
    description = "列出目录内容，支持递归、详情和文件过滤。"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "目录路径，默认为当前目录", "example": "src/"},
            "recursive": {"type": "boolean", "description": "是否递归列出子目录", "default": False, "example": True},
            "show_details": {"type": "boolean", "description": "显示文件大小和修改时间", "default": False},
            "pattern": {"type": "string", "description": "文件名 glob 过滤", "example": "*.py"},
            "max_depth": {"type": "integer", "description": "递归最大深度", "default": 3, "example": 2},
        },
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, path: str = ".", recursive: bool = False,
                show_details: bool = False, pattern: str = "",
                max_depth: int = 3) -> str:
        result = self.execute_structured(path, recursive, show_details, pattern, max_depth)
        return result.to_json()

    def execute_structured(self, path: str = ".", recursive: bool = False,
                           show_details: bool = False, pattern: str = "",
                           max_depth: int = 3) -> ToolResult:
        try:
            full_path = self.base_dir / Path(path)
            if not full_path.exists():
                return ToolResult.err(f"目录不存在：{path}")
            if not full_path.is_dir():
                return ToolResult.err(f"不是目录：{path}")

            dirs = []
            files = []
            lines = []

            def _format_size(size: int) -> str:
                if size < 1024:
                    return f"{size}B"
                elif size < 1024 * 1024:
                    return f"{size/1024:.1f}KB"
                elif size < 1024 * 1024 * 1024:
                    return f"{size/(1024*1024):.1f}MB"
                else:
                    return f"{size/(1024*1024*1024):.1f}GB"

            def _list_dir(current: Path, depth: int = 0):
                if depth > max_depth:
                    return
                try:
                    items = sorted(current.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
                except PermissionError:
                    lines.append("  " * depth + "⚠️ [无权限]")
                    return

                for item in items:
                    if item.name.startswith(".") and depth > 0:
                        continue
                    rel = str(item.relative_to(full_path))
                    if item.is_dir():
                        dirs.append(rel)
                        prefix = "📁"
                        detail = ""
                        if show_details:
                            try:
                                mtime = item.stat().st_mtime
                                from datetime import datetime as _dt
                                detail = f"  [{_dt.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M')}]"
                            except:
                                pass
                        lines.append("  " * depth + f"{prefix} {item.name}{detail}")
                        if recursive:
                            _list_dir(item, depth + 1)
                    else:
                        if pattern and not item.match(pattern):
                            continue
                        files.append(rel)
                        prefix = "📄"
                        detail = ""
                        if show_details:
                            try:
                                stat = item.stat()
                                size_str = _format_size(stat.st_size)
                                from datetime import datetime as _dt
                                mtime_str = _dt.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M')
                                detail = f"  [{size_str}, {mtime_str}]"
                            except:
                                pass
                        lines.append("  " * depth + f"{prefix} {item.name}{detail}")

            _list_dir(full_path)

            text = f"--- {path} ({len(dirs)} 目录, {len(files)} 文件) ---\n" + "\n".join(lines)
            return ToolResult.ok(
                data={"path": path, "directories": dirs, "files": files,
                      "total_dirs": len(dirs), "total_files": len(files)},
                message=text
            )
        except Exception as e:
            return ToolResult.err(f"列出目录失败：{e}")


class GrepTool(Tool):
    name = "grep"
    category = "search_code"
    keywords = ["grep", "search", "搜索", "查找", "代码"]
    model_hint = "在代码中搜索文本或正则表达式模式时使用。"
    description = "在文件中搜索文本或正则表达式，支持上下文显示。"
    parameters = {
        "type": "object",
        "properties": {
            "pattern": {"type": "string", "description": "【必填】要搜索的正则表达式或纯文本", "example": "def \\w+\\("},
            "path": {"type": "string", "description": "搜索目录（可选，默认整个项目）", "example": "src/"},
            "extension": {"type": "string", "description": "文件扩展名过滤", "example": ".py"},
            "ignore_case": {"type": "boolean", "description": "是否忽略大小写（默认 false）", "default": False},
            "context_lines": {"type": "integer", "description": "匹配行前后显示的上下文行数（默认 0）", "default": 0, "example": 3},
            "max_results": {"type": "integer", "description": "最大返回结果数量（默认 50）", "default": 50, "example": 20},
        },
        "required": ["pattern"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, pattern: str, path: str = ".", extension: str = "",
                ignore_case: bool = False, context_lines: int = 0,
                max_results: int = 50) -> str:
        result = self.execute_structured(pattern, path, extension, ignore_case, context_lines, max_results)
        return result.to_json()

    def execute_structured(self, pattern: str, path: str = ".", extension: str = "",
                           ignore_case: bool = False, context_lines: int = 0,
                           max_results: int = 50) -> ToolResult:
        try:
            full_path = self.base_dir / Path(path)
            if not full_path.exists():
                return ToolResult.err(f"路径不存在：{path}")

            flags = re.IGNORECASE if ignore_case else 0
            pattern_re = re.compile(pattern, flags)
            results = []
            match_count = 0
            file_count = 0

            def search_file(file_path: Path):
                nonlocal match_count, file_count
                try:
                    content = _read_file_with_encoding(file_path)
                    all_lines = content.splitlines()
                    found = False
                    for i, line in enumerate(all_lines):
                        if pattern_re.search(line):
                            if not found:
                                found = True
                                file_count += 1
                            if match_count < max_results:
                                if context_lines > 0:
                                    ctx_start = max(0, i - context_lines)
                                    ctx_end = min(len(all_lines), i + context_lines + 1)
                                    ctx_lines = []
                                    for j in range(ctx_start, ctx_end):
                                        marker = ">>>" if j == i else "   "
                                        ctx_lines.append(f"  {marker} {j+1:4d} | {all_lines[j]}")
                                    results.append(f"{file_path}:\n" + "\n".join(ctx_lines))
                                else:
                                    results.append(f"{file_path}:{i+1}: {line}")
                            match_count += 1
                            if match_count >= max_results + 1:
                                return True
                except Exception:
                    pass
                return False

            if full_path.is_file():
                search_file(full_path)
            else:
                for item in full_path.rglob("*" + extension if extension else "*"):
                    if item.is_file() and not item.name.startswith("."):
                        if search_file(item):
                            break

            if match_count == 0:
                return ToolResult.ok(
                    data={"pattern": pattern, "total_matches": 0, "total_files": 0, "results": []},
                    message=f"未找到匹配：{pattern}"
                )

            text = f"--- 搜索结果：{pattern} ---\n"
            text += "\n".join(results[:max_results])
            if match_count > max_results:
                text += f"\n\n... 还有 {match_count - max_results} 处匹配（共 {file_count} 个文件）"

            return ToolResult.ok(
                data={
                    "pattern": pattern,
                    "total_matches": match_count,
                    "total_files": file_count,
                    "showing": min(match_count, max_results),
                    "results": results[:max_results],
                },
                message=text
            )
        except Exception as e:
            return ToolResult.err(f"搜索失败：{e}")


class GlobTool(Tool):
    name = "glob"
    category = "search_code"
    keywords = ["glob", "find", "查找", "文件匹配", "pattern"]
    model_hint = "用文件扩展名或通配符模式查找文件路径时使用。"
    description = "用 glob 模式匹配查找文件路径。"
    parameters = {
        "type": "object",
        "properties": {
            "pattern": {"type": "string", "description": "glob 模式", "example": "**/*.py"},
            "path": {"type": "string", "description": "【必填】要匹配的搜索路径/目录", "example": "src/"},
            "show_details": {"type": "boolean", "description": "显示文件大小", "default": False},
        },
        "required": ["pattern"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, pattern: str, path: str = ".", show_details: bool = False) -> str:
        result = self.execute_structured(pattern, path, show_details)
        return result.to_json()

    def execute_structured(self, pattern: str, path: str = ".", show_details: bool = False) -> ToolResult:
        try:
            full_path = self.base_dir / Path(path)
            if not full_path.exists():
                return ToolResult.err(f"路径不存在：{path}")

            files = sorted(full_path.glob(pattern))
            if not files:
                return ToolResult.ok(
                    data={"pattern": pattern, "total": 0, "results": []},
                    message=f"未找到匹配：{pattern}"
                )

            def _fmt_size(size: int) -> str:
                if size < 1024:
                    return f"{size}B"
                elif size < 1024 * 1024:
                    return f"{size/1024:.1f}KB"
                else:
                    return f"{size/(1024*1024):.1f}MB"

            rel_files = []
            display_lines = []
            for f in files[:200]:
                try:
                    rel = str(f.relative_to(self.base_dir))
                except ValueError:
                    rel = str(f)
                rel_files.append(rel)
                if show_details and f.is_file():
                    try:
                        size = _fmt_size(f.stat().st_size)
                        display_lines.append(f"  {rel}  [{size}]")
                    except:
                        display_lines.append(f"  {rel}")
                else:
                    display_lines.append(f"  {rel}")

            text = f"--- 查找结果：{pattern} ({len(files)} 个) ---\n"
            text += "\n".join(display_lines)
            if len(files) > 200:
                text += f"\n\n... 还有 {len(files) - 200} 个文件"

            return ToolResult.ok(
                data={"pattern": pattern, "total": len(files), "results": rel_files[:200]},
                message=text
            )
        except Exception as e:
            return ToolResult.err(f"glob 查找失败：{e}")


class CalculateTool(Tool):
    name = "calculate"
    category = "math"
    keywords = ["calculate", "math", "计算", "数学", "表达式"]
    description = "数学计算，支持 Python 表达式和 numpy/scipy"
    parameters = {
        "type": "object",
        "properties": {
            "expression": {"type": "string", "description": "【必填】数学表达式，如 2+2, sin(pi/2), np.linalg.norm([3,4])。支持 Python math 和 numpy"},
            "timeout": {"type": "integer", "description": "计算超时秒数（默认 5，最大 30）"},
        },
        "required": ["expression"],
    }

    def execute(self, expression: str, timeout: int = 5) -> str:
        result = self.execute_structured(expression, timeout)
        return result.to_json()

    def execute_structured(self, expression: str, timeout: int = 5) -> ToolResult:
        try:
            import math
            safe_namespace = {k: v for k, v in math.__dict__.items() if not k.startswith("_")}
            safe_namespace["pi"] = math.pi
            safe_namespace["e"] = math.e

            try:
                import numpy as np
                safe_namespace["np"] = np
                safe_namespace["numpy"] = np
            except ImportError:
                pass

            try:
                import scipy
                safe_namespace["scipy"] = scipy
            except ImportError:
                pass

            import signal

            def _timeout_handler(signum, frame):
                raise TimeoutError(f"计算超时（{timeout}秒）")

            old_handler = None
            try:
                old_handler = signal.signal(signal.SIGALRM, _timeout_handler)
                signal.alarm(timeout)
            except (AttributeError, ValueError):
                pass

            try:
                result = eval(expression, {"__builtins__": {}}, safe_namespace)
            finally:
                try:
                    signal.alarm(0)
                    if old_handler is not None:
                        signal.signal(signal.SIGALRM, old_handler)
                except (AttributeError, ValueError):
                    pass

            try:
                result_str = str(result)
                if hasattr(result, '__len__') and not isinstance(result, str):
                    if len(result_str) > 500:
                        result_str = result_str[:500] + f"... (长度={len(result)})"
            except:
                result_str = str(result)

            return ToolResult.ok(
                data={"expression": expression, "result": result_str},
                message=f"{expression} = {result_str}"
            )
        except TimeoutError as e:
            return ToolResult.err(str(e))
        except Exception as e:
            return ToolResult.err(f"计算失败：{e}")


class WebSearchTool(Tool):
    name = "web_search"
    category = "web"
    keywords = ["web", "search", "搜索", "网络", "查询", "百度", "baidu", "博查", "bocha"]
    model_hint = "需要联网搜索获取最新信息时使用。支持博查AI搜索、百度AI搜索和DuckDuckGo。"
    description = "联网搜索查询。支持多搜索引擎（博查AI搜索/百度AI搜索/DuckDuckGo）、时间范围过滤、结果数量控制。"
    parameters = {
        "type": "object",
        "properties": {
            "query": {"type": "string", "description": "【必填】搜索关键词，支持多语言", "example": "Python 3.12 新特性"},
            "num_results": {"type": "integer", "description": "返回结果数量（默认 10）", "default": 10, "example": 5},
            "engine": {
                "type": "string",
                "description": "搜索引擎：auto(并发博查+DDG)/bocha(博查AI搜索)/baidu(百度AI搜索)/duckduckgo",
                "default": "auto",
                "enum": ["auto", "bocha", "baidu", "duckduckgo"],
            },
            "freshness": {
                "type": "string",
                "description": "时间范围过滤。博查：noLimit/oneDay/oneWeek/oneMonth/oneYear；百度：pd/pw/pm/py；或 YYYY-MM-DD..YYYY-MM-DD",
            },
        },
        "required": ["query"],
    }

    def execute(self, query: str, num_results: int = 10, engine: str = "auto", freshness: str = "") -> str:
        result = self.execute_structured(query, num_results, engine, freshness)
        return result.to_json()

    def execute_structured(self, query: str, num_results: int = 10, engine: str = "auto", freshness: str = "") -> ToolResult:
        if engine == "bocha":
            return self._search_bocha(query, num_results, freshness)
        elif engine == "baidu":
            return self._search_baidu(query, num_results, freshness)
        elif engine == "duckduckgo":
            return self._search_duckduckgo(query, num_results)
        else:
            # auto: 并发搜索两个引擎，合并去重
            return self._search_concurrent(query, num_results, freshness)

    def _search_concurrent(self, query: str, num_results: int, freshness: str) -> ToolResult:
        """并发搜索博查和 DuckDuckGo，合并去重结果"""
        import concurrent.futures

        bocha_ok = self._bocha_available()
        results_map = {}

        def _do_bocha():
            try:
                return self._search_bocha(query, num_results, freshness)
            except Exception as e:
                logger.debug(f"并发搜索-博查异常: {e}")
                return None

        def _do_ddg():
            try:
                return self._search_duckduckgo(query, num_results)
            except Exception as e:
                logger.debug(f"并发搜索-DDG异常: {e}")
                return None

        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as pool:
            futures = {}
            if bocha_ok:
                futures[pool.submit(_do_bocha)] = "bocha"
            futures[pool.submit(_do_ddg)] = "duckduckgo"

            try:
                for future in concurrent.futures.as_completed(futures, timeout=40):
                    engine_name = futures[future]
                    try:
                        result = future.result()
                        if result and result.success:
                            results_map[engine_name] = result
                    except Exception:
                        pass
            except TimeoutError:
                # 部分引擎超时，用已完成的继续
                logger.debug("并发搜索部分引擎超时，使用已完成的结果")

        # 合并去重
        seen_urls = set()
        merged = []
        engines_used = []

        # 博查结果优先（中文搜索质量更高，含摘要）
        for eng in ("bocha", "duckduckgo"):
            r = results_map.get(eng)
            if not r:
                continue
            engines_used.append(eng)
            for item in (r.data.get("results", []) if r.data else []):
                url = item.get("url", "")
                if url and url in seen_urls:
                    continue
                if url:
                    seen_urls.add(url)
                merged.append(item)

        if not merged:
            # 全部失败
            if results_map:
                # 至少有一个返回了但没结果
                first = list(results_map.values())[0]
                return first
            return ToolResult.err("所有搜索引擎均失败")

        # 截断到请求数量
        merged = merged[:num_results]

        lines = [f"--- 并发搜索：{query} ({'+'.join(engines_used)}) ---"]
        for i, item in enumerate(merged, 1):
            title = item.get("title", "")
            url_str = item.get("url", "")
            snippet = item.get("snippet", "")
            date = item.get("date", "")
            lines.append(f"{i}. {title}")
            if date:
                lines.append(f"   日期: {date}")
            lines.append(f"   {url_str}")
            if snippet:
                lines.append(f"   {snippet[:200]}")
            lines.append("")

        return ToolResult.ok(
            data={
                "query": query,
                "engine": "+".join(engines_used),
                "num_results": len(merged),
                "results": merged,
            },
            message="\n".join(lines)
        )

    def _baidu_available(self) -> bool:
        """检查百度搜索 API Key 是否可用"""
        import os
        _load_env_file()
        return bool(os.environ.get("BAIDU_API_KEY", ""))

    def _bocha_available(self) -> bool:
        """检查博查搜索 API Key 是否可用"""
        import os
        _load_env_file()
        return bool(os.environ.get("BOCHA_API_KEY", ""))

    @staticmethod
    def _map_freshness_to_bocha(freshness: str) -> str:
        """将通用 freshness 参数映射为博查 API 格式"""
        if not freshness:
            return "noLimit"
        mapping = {
            "pd": "oneDay", "24h": "oneDay", "oneDay": "oneDay",
            "pw": "oneWeek", "7d": "oneWeek", "oneWeek": "oneWeek",
            "pm": "oneMonth", "31d": "oneMonth", "oneMonth": "oneMonth",
            "py": "oneYear", "365d": "oneYear", "oneYear": "oneYear",
            "noLimit": "noLimit",
        }
        return mapping.get(freshness, "noLimit")

    def _search_bocha(self, query: str, count: int, freshness: str = "") -> ToolResult:
        """通过博查 AI Web Search API 搜索"""
        import urllib.request
        import urllib.error
        import os

        _load_env_file()
        api_key = os.environ.get("BOCHA_API_KEY", "")
        if not api_key:
            return self._search_duckduckgo(query, count)

        url = "https://api.bocha.cn/v1/web-search"

        # 构建请求体
        body = {
            "query": query,
            "count": min(count, 50),
            "freshness": self._map_freshness_to_bocha(freshness),
            "summary": True,
        }

        data = json.dumps(body, ensure_ascii=False).encode("utf-8")
        req = urllib.request.Request(
            url,
            data=data,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            method="POST",
        )

        try:
            with urllib.request.urlopen(req, timeout=30) as resp:
                resp_data = json.loads(resp.read().decode("utf-8"))

            # 提取搜索结果 — 响应格式兼容 Bing Search API
            web_pages = resp_data.get("data", {}).get("webPages", {})
            values = web_pages.get("value", [])
            lines = [f"--- 博查搜索：{query} ---"]
            search_results = []
            for i, item in enumerate(values[:count], 1):
                title = item.get("name", "")
                url_str = item.get("url", "")
                snippet = item.get("snippet", "")
                summary = item.get("summary", "")
                date = item.get("datePublished", "")
                site_name = item.get("siteName", "")
                # 优先使用 summary，其次 snippet
                content = summary or snippet
                lines.append(f"{i}. {title}")
                if site_name:
                    lines.append(f"   来源: {site_name}")
                if date:
                    lines.append(f"   日期: {date}")
                lines.append(f"   {url_str}")
                if content:
                    lines.append(f"   {content[:200]}")
                lines.append("")
                search_results.append({
                    "title": title,
                    "url": url_str,
                    "snippet": content[:300],
                    "date": date,
                    "site_name": site_name,
                })

            return ToolResult.ok(
                data={
                    "query": query,
                    "engine": "bocha",
                    "num_results": len(search_results),
                    "results": search_results,
                },
                message="\n".join(lines)
            )
        except urllib.error.HTTPError as e:
            error_body = ""
            try:
                error_body = e.read().decode("utf-8", errors="replace")
            except Exception:
                pass
            logger.warning(f"博查搜索失败 HTTP {e.code}: {error_body[:200]}，回退 DuckDuckGo")
            return self._search_duckduckgo(query, count)
        except Exception as e:
            logger.warning(f"博查搜索异常: {e}，回退 DuckDuckGo")
            return self._search_duckduckgo(query, count)

    def _search_baidu(self, query: str, count: int, freshness: str) -> ToolResult:
        """通过百度 AI Search API 搜索"""
        import urllib.request
        import urllib.error
        import os

        _load_env_file()
        api_key = os.environ.get("BAIDU_API_KEY", "")
        if not api_key:
            # 回退到 DuckDuckGo
            return self._search_duckduckgo(query, count)

        url = "https://qianfan.baidubce.com/v2/ai_search/web_search"

        # 构建请求体
        body = MCPServiceTool._transform_baidu_search({
            "query": query,
            "count": count,
            "freshness": freshness,
        })

        data = json.dumps(body, ensure_ascii=False).encode("utf-8")
        req = urllib.request.Request(
            url,
            data=data,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "X-Appbuilder-From": "openclaw",
            },
            method="POST",
        )

        try:
            with urllib.request.urlopen(req, timeout=30) as resp:
                result = json.loads(resp.read().decode("utf-8"))

            # 提取搜索结果
            refs = result.get("references", [])
            lines = [f"--- 百度搜索：{query} ---"]
            search_results = []
            for i, ref in enumerate(refs[:count], 1):
                title = ref.get("title", "")
                url_str = ref.get("url", "")
                content = ref.get("content", ref.get("snippet", ""))
                date = ref.get("date", "")
                lines.append(f"{i}. {title}")
                if date:
                    lines.append(f"   日期: {date}")
                lines.append(f"   {url_str}")
                if content:
                    lines.append(f"   {content[:200]}")
                lines.append("")
                search_results.append({
                    "title": title,
                    "url": url_str,
                    "snippet": content[:300],
                    "date": date,
                })

            return ToolResult.ok(
                data={
                    "query": query,
                    "engine": "baidu",
                    "num_results": len(search_results),
                    "results": search_results,
                },
                message="\n".join(lines)
            )
        except urllib.error.HTTPError as e:
            error_body = ""
            try:
                error_body = e.read().decode("utf-8", errors="replace")
            except Exception:
                pass
            # 百度失败时回退 DuckDuckGo
            logger.warning(f"百度搜索失败 HTTP {e.code}，回退 DuckDuckGo")
            return self._search_duckduckgo(query, count)
        except Exception as e:
            logger.warning(f"百度搜索异常: {e}，回退 DuckDuckGo")
            return self._search_duckduckgo(query, count)

    def _search_duckduckgo(self, query: str, num_results: int) -> ToolResult:
        """通过 DuckDuckGo 搜索"""
        try:
            try:
                from ddgs import DDGS
            except ImportError:
                from duckduckgo_search import DDGS
            results = list(DDGS(timeout=30).text(query, max_results=num_results))
            lines = [f"--- DuckDuckGo 搜索：{query} ---"]
            search_results = []
            for i, r in enumerate(results, 1):
                lines.append(f"{i}. {r.get('title', '')}")
                lines.append(f"   {r.get('href', '')}")
                lines.append(f"   {r.get('body', '')[:150]}...")
                lines.append("")
                search_results.append({
                    "title": r.get("title", ""),
                    "url": r.get("href", ""),
                    "snippet": r.get("body", "")[:300],
                })

            return ToolResult.ok(
                data={
                    "query": query,
                    "engine": "duckduckgo",
                    "num_results": len(results),
                    "results": search_results,
                },
                message="\n".join(lines)
            )
        except ImportError:
            return ToolResult.err(
                "百度搜索不可用（未配置 API Key），DuckDuckGo 也未安装。\n"
                "请设置 BAIDU_API_KEY 或运行：pip install duckduckgo-search"
            )
        except Exception as e:
            return ToolResult.err(f"搜索失败：{e}")


class FetchUrlTool(Tool):
    name = "fetch_url"
    category = "web"
    keywords = ["fetch", "url", "网页", "获取", "下载"]
    model_hint = "需要获取指定网页内容或抓取网页数据时使用。"
    description = "获取指定 URL 的网页文本内容。"
    parameters = {
        "type": "object",
        "properties": {
            "url": {"type": "string", "description": "网页 URL", "example": "https://example.com"},
            "max_chars": {"type": "integer", "description": "最大返回字符数（默认 10000）", "default": 10000, "example": 5000},
        },
        "required": ["url"],
    }

    def execute(self, url: str, max_chars: int = 10000) -> str:
        result = self.execute_structured(url, max_chars)
        return result.to_json()

    def execute_structured(self, url: str, max_chars: int = 10000) -> ToolResult:
        try:
            import urllib.request
            import re as _re

            req = urllib.request.Request(url, headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            })
            with urllib.request.urlopen(req, timeout=15) as resp:
                html = resp.read().decode("utf-8", errors="ignore")

            text = _re.sub(r'<script[^>]*>.*?</script>', '', html, flags=_re.DOTALL | _re.IGNORECASE)
            text = _re.sub(r'<style[^>]*>.*?</style>', '', text, flags=_re.DOTALL | _re.IGNORECASE)
            text = _re.sub(r'<br\s*/?>', '\n', text, flags=_re.IGNORECASE)
            text = _re.sub(r'<p[^>]*>', '\n', text, flags=_re.IGNORECASE)
            text = _re.sub(r'<[^>]+>', '', text)
            text = _re.sub(r'&nbsp;', ' ', text)
            text = _re.sub(r'&amp;', '&', text)
            text = _re.sub(r'&lt;', '<', text)
            text = _re.sub(r'&gt;', '>', text)
            text = _re.sub(r'&#\d+;', '', text)
            text = _re.sub(r'\n\s*\n', '\n\n', text)
            text = text.strip()

            original_len = len(text)
            truncated = original_len > max_chars
            if truncated:
                text = text[:max_chars] + f"\n\n... [已截断，原始 {original_len} 字符]"

            return ToolResult.ok(
                data={"url": url, "content": text, "original_length": original_len, "truncated": truncated},
                message=f"--- {url} ---\n{text}"
            )
        except Exception as e:
            return ToolResult.err(f"获取网页失败：{e}")


class AnalyzePaperTool(Tool):
    name = "analyze_paper"
    category = "research"
    keywords = ["paper", "analyze", "论文", "分析", "研究"]
    model_hint = "分析 PDF/DOCX/MD 格式的科研论文内容时使用。"
    description = "读取并分析论文文件（PDF/DOCX/MD），可选提取元数据。"
    parameters = {
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "论文文件路径", "example": "papers/machine_learning.pdf"},
            "start_page": {"type": "integer", "description": "起始页码（仅 PDF）", "example": 1},
            "end_page": {"type": "integer", "description": "结束页码（仅 PDF）", "example": 5},
            "extract_metadata": {"type": "boolean", "description": "提取元数据（标题/作者等）", "default": False},
        },
        "required": ["file_path"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, file_path: str, start_page: Optional[int] = None,
                end_page: Optional[int] = None, extract_metadata: bool = False) -> str:
        result = self.execute_structured(file_path, start_page, end_page, extract_metadata)
        return result.to_json()

    def execute_structured(self, file_path: str, start_page: Optional[int] = None,
                           end_page: Optional[int] = None, extract_metadata: bool = False) -> ToolResult:
        try:
            full_path = self.base_dir / Path(file_path)
            if not full_path.exists():
                return ToolResult.err(f"文件不存在：{file_path}")

            metadata = {}
            if extract_metadata:
                metadata = self._extract_metadata(full_path)

            text = self._extract_text(full_path, start_page, end_page)
            if not text:
                return ToolResult.err("无法提取文本内容")

            truncated = len(text) > self.max_result_chars
            preview = text[:self.max_result_chars]

            data = {
                "file_path": file_path,
                "original_length": len(text),
                "preview": preview,
                "truncated": truncated,
            }
            if metadata:
                data["metadata"] = metadata

            header = f"--- 分析论文：{file_path} ---"
            if metadata:
                header += "\n" + "\n".join(f"  {k}: {v}" for k, v in metadata.items() if v)
                header += "\n"

            return ToolResult.ok(data=data, message=f"{header}\n{preview}")
        except Exception as e:
            return ToolResult.err(f"分析论文失败：{e}")

    def _extract_text(self, path: Path, start_page=None, end_page=None) -> str:
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                try:
                    import PyPDF2
                    with open(path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        pages = reader.pages
                        total = len(pages)
                        s = max(0, (start_page or 1) - 1)
                        e = min(total, end_page or total)
                        return "\n".join([pages[i].extract_text() or "" for i in range(s, e)])
                except ImportError:
                    return "提示：需要安装 PyPDF2 才能读取 PDF\n运行：pip install PyPDF2"
            elif ext in [".docx", ".doc"]:
                try:
                    from docx import Document
                    doc = Document(path)
                    return "\n".join([p.text for p in doc.paragraphs])
                except ImportError:
                    return "提示：需要安装 python-docx 才能读取 DOCX\n运行：pip install python-docx"
            else:
                return _read_file_with_encoding(path)
        except Exception as e:
            return f"提取文本失败：{e}"

    def _extract_metadata(self, path: Path) -> dict:
        meta = {}
        ext = path.suffix.lower()
        if ext == ".pdf":
            try:
                import PyPDF2
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    info = reader.metadata
                    if info:
                        meta["title"] = info.get("/Title", "")
                        meta["author"] = info.get("/Author", "")
                        meta["subject"] = info.get("/Subject", "")
                        meta["creator"] = info.get("/Creator", "")
                    meta["total_pages"] = len(reader.pages)
            except:
                pass
        elif ext in [".docx", ".doc"]:
            try:
                from docx import Document
                doc = Document(path)
                props = doc.core_properties
                meta["title"] = props.title or ""
                meta["author"] = props.author or ""
                meta["created"] = str(props.created) if props.created else ""
                meta["modified"] = str(props.modified) if props.modified else ""
            except:
                pass
        meta["file_size"] = path.stat().st_size
        return meta


class OpenFileTool(Tool):
    name = "open_file"
    category = "file_io"
    keywords = ["open", "file", "打开", "文件"]
    description = "用系统默认程序打开文件（如 PDF 用阅读器、图片用查看器等）。仅打开文件，不返回内容。"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "【必填】要打开的文件路径（相对于工作目录或被打开前的绝对路径）"},
        },
        "required": ["path"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()

    def execute(self, path: str) -> str:
        result = self.execute_structured(path)
        return result.to_json()

    def execute_structured(self, path: str) -> ToolResult:
        import subprocess
        import platform

        target = Path(path)
        if not target.is_absolute():
            target = self.base_dir / target

        if not target.exists():
            return ToolResult.err(f"文件不存在: {target}")

        try:
            system = platform.system()
            if system == "Windows":
                os.startfile(str(target))
            elif system == "Darwin":
                subprocess.run(["open", str(target)], check=True)
            else:
                subprocess.run(["xdg-open", str(target)], check=True)

            return ToolResult.ok(
                data={"path": str(target), "opened": True},
                message=f"✅ 已用系统默认方式打开: {target.name}"
            )
        except Exception as e:
            return ToolResult.err(f"打开文件失败: {e}")


class TerminalTool(Tool):
    name = "terminal_open"
    category = "system"
    keywords = ["terminal", "shell", "终端", "命令行"]
    description = "打开新的终端窗口（可选：在指定目录打开，或执行命令后保持打开）"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "工作目录（可选，默认项目根目录）"},
            "command": {"type": "string", "description": "要在终端中执行的命令（可选，留空则仅打开终端）"},
            "keep_open": {"type": "boolean", "description": "命令执行后是否保持终端窗口不关闭（默认 true）"},
        },
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()

    def execute(self, path: Optional[str] = None, command: Optional[str] = None, keep_open: bool = True) -> str:
        result = self.execute_structured(path, command, keep_open)
        return result.to_json()

    def execute_structured(self, path: Optional[str] = None, command: Optional[str] = None, keep_open: bool = True) -> ToolResult:
        import subprocess
        import platform
        import sys

        working_dir = self.base_dir
        if path:
            p = Path(path)
            if p.is_absolute():
                working_dir = p
            else:
                working_dir = self.base_dir / p
        
        working_dir.mkdir(parents=True, exist_ok=True)

        try:
            system = platform.system()
            
            if system == "Windows":
                cmd_parts = ["cmd", "/k"] if keep_open else ["cmd", "/c"]
                if command:
                    cmd_parts.append(command)
                subprocess.Popen(cmd_parts, cwd=str(working_dir), creationflags=subprocess.CREATE_NEW_CONSOLE)
                
            elif system == "Darwin":
                script = f'cd "{working_dir}"'
                if command:
                    script += f'; {command}'
                if keep_open:
                    script += '; exec bash'
                applescript = f'tell application "Terminal" to do script "{script}"'
                subprocess.run(["osascript", "-e", applescript], check=True)
                
            else:
                shell = os.environ.get("SHELL", "bash")
                if command:
                    if keep_open:
                        script = f"{command}; exec {shell}"
                    else:
                        script = command
                    subprocess.Popen([shell, "-c", script], cwd=str(working_dir))
                else:
                    subprocess.Popen([shell], cwd=str(working_dir))

            return ToolResult.ok(
                data={"path": str(working_dir), "command": command, "keep_open": keep_open},
                message=f"✅ 已打开终端" + (f" 并执行: {command}" if command else "")
            )
        except Exception as e:
            return ToolResult.err(f"打开终端失败: {e}")


class DiffFilesTool(Tool):
    name = "diff_files"
    category = "file_io"
    keywords = ["diff", "compare", "比较", "差异"]
    description = "比较两个文件的差异"
    parameters = {
        "type": "object",
        "properties": {
            "path_a": {"type": "string", "description": "【必填】要比较的第一个文件路径"},
            "path_b": {"type": "string", "description": "【必填】要比较的第二个文件路径"},
            "context_lines": {"type": "integer", "description": "差异上下文行数（默认 3，显示更多上下文请增大数值）"},
        },
        "required": ["path_a", "path_b"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, path_a: str, path_b: str, context_lines: int = 3) -> str:
        result = self.execute_structured(path_a, path_b, context_lines)
        return result.to_json()

    def execute_structured(self, path_a: str, path_b: str, context_lines: int = 3) -> ToolResult:
        try:
            import difflib
            full_a = self.base_dir / Path(path_a)
            full_b = self.base_dir / Path(path_b)
            if not full_a.exists():
                return ToolResult.err(f"文件不存在：{path_a}")
            if not full_b.exists():
                return ToolResult.err(f"文件不存在：{path_b}")

            lines_a = _read_file_with_encoding(full_a).splitlines(keepends=True)
            lines_b = _read_file_with_encoding(full_b).splitlines(keepends=True)

            diff = difflib.unified_diff(lines_a, lines_b, fromfile=path_a, tofile=path_b, n=context_lines)
            diff_text = "".join(diff)

            if not diff_text:
                return ToolResult.ok(
                    data={"path_a": path_a, "path_b": path_b, "has_diff": False},
                    message=f"✅ {path_a} 和 {path_b} 内容完全相同"
                )

            if len(diff_text) > self.max_result_chars:
                diff_text = diff_text[:self.max_result_chars] + "\n... [diff 已截断]"

            return ToolResult.ok(
                data={"path_a": path_a, "path_b": path_b, "has_diff": True},
                message=f"--- 差异：{path_a} vs {path_b} ---\n{diff_text}"
            )
        except Exception as e:
            return ToolResult.err(f"比较文件失败：{e}")


class MoveFileTool(Tool):
    name = "move_file"
    category = "file_io"
    keywords = ["move", "rename", "移动", "重命名"]
    description = "移动或重命名文件"
    parameters = {
        "type": "object",
        "properties": {
            "source": {"type": "string", "description": "【必填】源文件路径（要移动/重命名的文件）"},
            "destination": {"type": "string", "description": "【必填】目标文件路径（新位置或新文件名）"},
        },
        "required": ["source", "destination"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, source: str, destination: str) -> str:
        result = self.execute_structured(source, destination)
        return result.to_json()

    def execute_structured(self, source: str, destination: str) -> ToolResult:
        try:
            import shutil
            src = self.base_dir / Path(source)
            dst = self.base_dir / Path(destination)
            if not src.exists():
                return ToolResult.err(f"源文件不存在：{source}")
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.move(str(src), str(dst))
            return ToolResult.ok(
                data={"source": source, "destination": destination},
                message=f"✅ 已移动: {source} → {destination}"
            )
        except Exception as e:
            return ToolResult.err(f"移动文件失败：{e}")


class CopyFileTool(Tool):
    name = "copy_file"
    category = "file_io"
    keywords = ["copy", "复制", "拷贝"]
    description = "复制文件"
    parameters = {
        "type": "object",
        "properties": {
            "source": {"type": "string", "description": "【必填】源文件路径（要复制的文件）"},
            "destination": {"type": "string", "description": "【必填】目标文件路径（副本存放位置）"},
        },
        "required": ["source", "destination"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, source: str, destination: str) -> str:
        result = self.execute_structured(source, destination)
        return result.to_json()

    def execute_structured(self, source: str, destination: str) -> ToolResult:
        try:
            import shutil
            src = self.base_dir / Path(source)
            dst = self.base_dir / Path(destination)
            if not src.exists():
                return ToolResult.err(f"源文件不存在：{source}")
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(str(src), str(dst))
            return ToolResult.ok(
                data={"source": source, "destination": destination},
                message=f"✅ 已复制: {source} → {destination}"
            )
        except Exception as e:
            return ToolResult.err(f"复制文件失败：{e}")


class FileInfoTool(Tool):
    name = "file_info"
    category = "file_io"
    keywords = ["info", "file", "文件信息", "属性"]
    description = "获取文件元数据（大小、修改时间、编码检测等）"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "【必填】要查看信息的文件路径（相对于工作目录）"},
        },
        "required": ["path"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, path: str) -> str:
        result = self.execute_structured(path)
        return result.to_json()

    def execute_structured(self, path: str) -> ToolResult:
        try:
            full_path = self.base_dir / Path(path)
            if not full_path.exists():
                return ToolResult.err(f"文件不存在：{path}")

            stat = full_path.stat()
            from datetime import datetime as _dt

            info = {
                "path": str(full_path),
                "name": full_path.name,
                "suffix": full_path.suffix,
                "size_bytes": stat.st_size,
                "size_readable": self._fmt_size(stat.st_size),
                "is_file": full_path.is_file(),
                "is_dir": full_path.is_dir(),
                "created": _dt.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M:%S"),
                "modified": _dt.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                "accessed": _dt.fromtimestamp(stat.st_atime).strftime("%Y-%m-%d %H:%M:%S"),
            }

            if full_path.is_file():
                try:
                    content = full_path.read_bytes()[:4096]
                    info["encoding"] = self._detect_encoding(content)
                    info["line_count"] = len(_read_file_with_encoding(full_path).splitlines())
                except:
                    pass

            lines = [f"--- 文件信息：{path} ---"]
            for k, v in info.items():
                lines.append(f"  {k}: {v}")

            return ToolResult.ok(data=info, message="\n".join(lines))
        except Exception as e:
            return ToolResult.err(f"获取文件信息失败：{e}")

    def _fmt_size(self, size: int) -> str:
        if size < 1024:
            return f"{size} B"
        elif size < 1024 * 1024:
            return f"{size/1024:.1f} KB"
        elif size < 1024 * 1024 * 1024:
            return f"{size/(1024*1024):.1f} MB"
        else:
            return f"{size/(1024*1024*1024):.1f} GB"

    def _detect_encoding(self, raw: bytes) -> str:
        try:
            import chardet
            result = chardet.detect(raw)
            return result.get("encoding", "unknown")
        except ImportError:
            pass
        if raw.startswith(b'\xef\xbb\xbf'):
            return "utf-8-sig"
        if raw.startswith(b'\xff\xfe'):
            return "utf-16-le"
        if raw.startswith(b'\xfe\xff'):
            return "utf-16-be"
        try:
            raw.decode("utf-8")
            return "utf-8"
        except:
            pass
        return "ascii/latin-1 (猜测)"


class CLITool(Tool):
    name = "cli_execute"
    category = "system"
    keywords = ["cli", "execute", "命令", "执行", "shell"]
    model_hint = "执行安全的、非破坏性的命令行命令时使用。"
    description = "执行命令行命令。"
    parameters = {
        "type": "object",
        "properties": {
            "command": {"type": "string", "description": "要执行的命令", "example": "ls -la"},
            "timeout": {"type": "integer", "description": "命令执行超时秒数（默认 30）", "default": 30, "example": 10},
            "cwd": {"type": "string", "description": "命令执行的工作目录（可选，默认项目根目录）", "example": "src/"},
            "open_terminal": {"type": "boolean", "description": "在新终端窗口执行", "default": False},
            "keep_terminal_open": {"type": "boolean", "description": "执行后保持终端打开", "default": True},
        },
        "required": ["command"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, command: str, timeout: int = 30, cwd: Optional[str] = None, open_terminal: bool = False, keep_terminal_open: bool = True) -> str:
        result = self.execute_structured(command, timeout, cwd, open_terminal, keep_terminal_open)
        return result.to_json()

    def execute_structured(self, command: str, timeout: int = 30, cwd: Optional[str] = None, open_terminal: bool = False, keep_terminal_open: bool = True) -> ToolResult:
        try:
            if open_terminal:
                return self._execute_in_terminal(command, cwd, keep_terminal_open)
            
            import subprocess
            import platform
            import shlex

            working_dir = cwd or str(self.base_dir)
            
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='replace',
                timeout=timeout,
                cwd=working_dir,
                shell=True
            )

            output = []
            if result.stdout:
                output.append(f"--- stdout ---\n{result.stdout}")
            if result.stderr:
                output.append(f"--- stderr ---\n{result.stderr}")

            output_text = "\n".join(output) if output else "(无输出)"

            if len(output_text) > self.max_result_chars:
                output_text = output_text[:self.max_result_chars] + "\n... [输出已截断]"

            data = {
                "command": command,
                "returncode": result.returncode,
                "stdout": result.stdout,
                "stderr": result.stderr,
            }

            if result.returncode == 0:
                return ToolResult.ok(
                    data=data,
                    message=f"✅ 命令执行成功 (返回码: 0)\n\n{output_text}"
                )
            else:
                return ToolResult.ok(
                    data=data,
                    message=f"⚠️ 命令执行完成 (返回码: {result.returncode})\n\n{output_text}"
                )
        except subprocess.TimeoutExpired:
            return ToolResult.err(f"命令执行超时 ({timeout}秒)")
        except Exception as e:
            return ToolResult.err(f"执行命令失败：{e}")

    def _execute_in_terminal(self, command: str, cwd: Optional[str], keep_open: bool) -> ToolResult:
        import subprocess
        import platform

        working_dir = self.base_dir
        if cwd:
            p = Path(cwd)
            if p.is_absolute():
                working_dir = p
            else:
                working_dir = self.base_dir / p
        
        working_dir.mkdir(parents=True, exist_ok=True)

        try:
            system = platform.system()
            
            if system == "Windows":
                cmd_parts = ["cmd", "/k"] if keep_open else ["cmd", "/c"]
                cmd_parts.append(command)
                subprocess.Popen(cmd_parts, cwd=str(working_dir), creationflags=subprocess.CREATE_NEW_CONSOLE)
                
            elif system == "Darwin":
                script = f'cd "{working_dir}"; {command}'
                if keep_open:
                    script += '; exec bash'
                applescript = f'tell application "Terminal" to do script "{script}"'
                subprocess.run(["osascript", "-e", applescript], check=True)
                
            else:
                shell = os.environ.get("SHELL", "bash")
                if keep_open:
                    script = f"{command}; exec {shell}"
                else:
                    script = command
                subprocess.Popen([shell, "-c", script], cwd=str(working_dir))

            return ToolResult.ok(
                data={"command": command, "in_terminal": True, "keep_open": keep_open},
                message=f"🖥️ 已在新终端中执行命令: {command}"
            )
        except Exception as e:
            return ToolResult.err(f"在终端执行失败: {e}")


class ConfigTool(Tool):
    name = "config_manage"
    category = "system"
    keywords = ["config", "设置", "配置", "管理", "ws2", "datahub"]
    model_hint = "查看/修改系统配置、模型提供商、WS2 课程和数据枢纽时使用。"
    description = "管理系统配置、WS2 系统状态、数据枢纽。"
    parameters = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "description": "操作类型：list_providers=列出, get_settings=读取, update_setting=修改, list_skills=列技能, ws2_stats=WS2统计",
                "enum": ["list_providers", "get_settings", "update_setting", "list_skills", "ws2_stats", "ws2_list", "ws2_reload", "hub_stats", "hub_list"],
                "example": "list_providers"
            },
            "key": {"type": "string", "description": "[update_setting] 设置键名", "example": "theme"},
            "value": {"type": "string", "description": "[update_setting] 设置值", "example": "dark"},
            "filter": {"type": "string", "description": "[ws2_list/hub_list] 过滤条件", "example": "机器学习"},
        },
        "required": ["action"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, action: str, key: Optional[str] = None, value: Optional[str] = None, filter: Optional[str] = None) -> str:
        result = self.execute_structured(action, key, value, filter)
        return result.to_json()

    def execute_structured(self, action: str, key: Optional[str] = None, value: Optional[str] = None, filter: Optional[str] = None) -> ToolResult:
        try:
            if action in ("ws2_stats", "ws2_list", "ws2_reload"):
                return self._handle_ws2(action, filter)
            elif action in ("hub_stats", "hub_list"):
                return self._handle_datahub(action, filter)
            else:
                return self._handle_config(action, key, value)
        except Exception as e:
            return ToolResult.err(f"配置管理失败：{e}")

    def _handle_ws2(self, action: str, filter: Optional[str] = None) -> ToolResult:
        try:
            from ws2_data_hub import get_data_hub
            hub = get_data_hub()
            if hub is None:
                return ToolResult.err("数据枢纽未初始化")

            if action == "ws2_stats":
                stats = hub.get_statistics()
                lines = ["=== WS2 系统统计 ==="]
                lines.append(f"课程总数: {stats.get('total_courses', 0)}")
                lines.append(f"课时总数: {stats.get('total_lessons', 0)}")
                lines.append(f"已完成课时: {stats.get('completed_lessons', 0)}")
                lines.append(f"总学时: {stats.get('total_hours', 0)}")
                lines.append(f"数据源数量: {len(hub.get_all_sources())}")
                return ToolResult.ok(data=stats, message="\n".join(lines))

            elif action == "ws2_list":
                courses = hub.list_courses(filter_keyword=filter) if filter else hub.list_courses()
                lines = ["=== WS2 课程列表 ==="]
                for c in courses[:20]:
                    lines.append(f"- {c.get('title', '未知')}: {c.get('lesson_count', 0)}课时")
                if len(courses) > 20:
                    lines.append(f"... 还有 {len(courses) - 20} 个课程")
                return ToolResult.ok(
                    data={"courses": courses[:20], "total": len(courses)},
                    message="\n".join(lines)
                )

            elif action == "ws2_reload":
                hub.reload_all_sources()
                return ToolResult.ok(message="✅ WS2 数据源已重新加载")

        except ImportError:
            return ToolResult.err("ws2_data_hub 模块不可用")
        except Exception as e:
            return ToolResult.err(f"WS2 操作失败：{e}")

    def _handle_datahub(self, action: str, filter: Optional[str] = None) -> ToolResult:
        try:
            from ws2_data_hub import get_data_hub
            hub = get_data_hub()
            if hub is None:
                return ToolResult.err("数据枢纽未初始化")

            if action == "hub_stats":
                stats = hub.get_statistics()
                lines = ["=== 数据枢纽统计 ==="]
                lines.append(f"总项目数: {stats.get('total_items', 0)}")
                lines.append(f"未读项目: {stats.get('unread_items', 0)}")
                lines.append(f"星标项目: {stats.get('starred_items', 0)}")
                sources = hub.get_all_sources()
                lines.append(f"数据源: {len(sources)}")
                for s in sources[:5]:
                    lines.append(f"  - {s.get('name', '未知')}: {s.get('item_count', 0)}项")
                return ToolResult.ok(data=stats, message="\n".join(lines))

            elif action == "hub_list":
                items = hub.query_items(search=filter, limit=20) if filter else hub.query_items(limit=20)
                lines = ["=== 数据枢纽项目 ==="]
                for item in items:
                    star = "⭐" if item.get("is_starred") else ""
                    lines.append(f"{star} {item.get('title', '未知')} [{item.get('item_type', '?')}]")
                if not items:
                    lines.append("(空)")
                return ToolResult.ok(
                    data={"items": items, "count": len(items)},
                    message="\n".join(lines)
                )

        except ImportError:
            return ToolResult.err("ws2_data_hub 模块不可用")
        except Exception as e:
            return ToolResult.err(f"数据枢纽操作失败：{e}")

    def _handle_config(self, action: str, key: Optional[str] = None, value: Optional[str] = None) -> ToolResult:
        from .config import get_config_manager
        config_mgr = get_config_manager()

        if action == "list_providers":
            providers = config_mgr.get_all_provider_configs()
            lines = ["--- 模型提供商配置 ---"]
            for p in providers:
                status = "✅" if p.enabled else "❌"
                lines.append(f"{status} {p.name} ({p.provider.value}) - 模型: {p.model}")
            return ToolResult.ok(
                data={"providers": [{"name": p.name, "provider": p.provider.value, "model": p.model, "enabled": p.enabled} for p in providers]},
                message="\n".join(lines)
            )
        
        elif action == "get_settings":
            settings = config_mgr.settings
            lines = ["--- 系统设置 ---"]
            for k, v in settings.items():
                lines.append(f"  {k}: {v}")
            return ToolResult.ok(
                data={"settings": settings},
                message="\n".join(lines) if settings else "暂无设置"
            )
        
        elif action == "update_setting":
            if key is None:
                return ToolResult.err("update_setting 需要 key 参数")
            config_mgr.set_setting(key, value or "")
            return ToolResult.ok(
                data={"key": key, "value": value},
                message=f"✅ 已设置 {key} = {value}"
            )
        
        elif action == "list_skills":
            skills = config_mgr.skill_configs
            lines = ["--- 技能配置 ---"]
            for s in skills:
                status = "✅" if s.enabled else "❌"
                lines.append(f"{status} {s.name} - {s.description}")
            return ToolResult.ok(
                data={"skills": [{"name": s.name, "description": s.description, "enabled": s.enabled} for s in skills]},
                message="\n".join(lines)
            )
        
        else:
            return ToolResult.err(f"未知操作：{action}")


class SkillTool(Tool):
    name = "skill_manager"
    category = "system"
    keywords = ["skill", "技能", "管理", "加载"]
    model_hint = "列出、查看或执行已注册的技能时使用。"
    description = "管理技能：列出可用技能、获取详情、执行技能、执行脚本。"
    parameters = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "description": "操作类型：list=列出, get_details=获取详情, execute=执行, execute_script=执行脚本",
                "enum": ["list", "get_details", "execute", "execute_script"],
                "example": "list"
            },
            "skill_name": {"type": "string", "description": "[get_details/execute/execute_script] 技能名称", "example": "analyze_paper"},
            "parameters": {"type": "object", "description": "[execute/execute_script] 技能执行参数"},
            "script_name": {"type": "string", "description": "[execute_script] 脚本名称（如 'main'），对应 SKILL.md 中 scripts 字段的键名", "example": "main"},
        },
        "required": ["action"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()
        self._skill_registry: Dict[str, Callable] = {}

    def execute(self, action: str, skill_name: Optional[str] = None, parameters: Optional[Dict[str, Any]] = None, script_name: Optional[str] = None) -> str:
        result = self.execute_structured(action, skill_name, parameters, script_name)
        return result.to_json()

    def execute_structured(self, action: str, skill_name: Optional[str] = None, parameters: Optional[Dict[str, Any]] = None, script_name: Optional[str] = None) -> ToolResult:
        try:
            from .config import get_config_manager
            config_mgr = get_config_manager()

            if action == "list":
                enabled_skills = config_mgr.get_enabled_skills()
                disabled_skills = [s for s in config_mgr.skill_configs.values() if not s.enabled]
                
                lines = ["=== 已启用的技能 ==="]
                for s in enabled_skills:
                    lines.append(f"✅ {s.name}: {s.description}")
                    if s.parameters:
                        lines.append(f"   参数: {json.dumps(s.parameters, ensure_ascii=False)}")
                
                if disabled_skills:
                    lines.append("\n=== 已禁用的技能 ===")
                    for s in disabled_skills:
                        lines.append(f"❌ {s.name}: {s.description}")
                
                data = {
                    "enabled": [{"name": s.name, "description": s.description, "type": s.type} for s in enabled_skills],
                    "disabled": [{"name": s.name, "description": s.description, "type": s.type} for s in disabled_skills],
                }
                
                return ToolResult.ok(
                    data=data,
                    message="\n".join(lines)
                )
            
            elif action == "get_details":
                if not skill_name:
                    return ToolResult.err("get_details 需要 skill_name 参数")
                
                skill_config = config_mgr.skill_configs.get(skill_name)
                if not skill_config:
                    return ToolResult.err(f"技能 '{skill_name}' 不存在")
                
                details = {
                    "name": skill_config.name,
                    "description": skill_config.description,
                    "type": skill_config.type,
                    "enabled": skill_config.enabled,
                    "parameters": skill_config.parameters,
                    "dependencies": skill_config.dependencies,
                    "handler_path": skill_config.handler_path,
                    "metadata": skill_config.metadata,
                }
                
                lines = [f"=== 技能详情: {skill_name} ==="]
                lines.append(f"描述: {skill_config.description}")
                lines.append(f"类型: {skill_config.type}")
                lines.append(f"状态: {'启用' if skill_config.enabled else '禁用'}")
                if skill_config.parameters:
                    lines.append(f"参数定义: {json.dumps(skill_config.parameters, ensure_ascii=False, indent=2)}")
                if skill_config.dependencies:
                    lines.append(f"依赖: {', '.join(skill_config.dependencies)}")
                if skill_config.handler_path:
                    lines.append(f"处理器路径: {skill_config.handler_path}")
                
                return ToolResult.ok(
                    data=details,
                    message="\n".join(lines)
                )
            
            elif action == "execute":
                if not skill_name:
                    return ToolResult.err("execute 需要 skill_name 参数")
                
                skill_config = config_mgr.skill_configs.get(skill_name)
                if not skill_config:
                    return ToolResult.err(f"技能 '{skill_name}' 不存在")
                
                if not skill_config.enabled:
                    return ToolResult.err(f"技能 '{skill_name}' 已禁用")
                
                result = self._execute_skill(skill_config, parameters or {})
                
                if result["success"]:
                    return ToolResult.ok(
                        data=result["data"],
                        message=result["message"]
                    )
                else:
                    return ToolResult.err(result["message"])
            
            elif action == "execute_script":
                if not skill_name:
                    return ToolResult.err("execute_script 需要 skill_name 参数")

                return self._execute_skill_script(skill_name, script_name or "main", parameters or {})

            else:
                return ToolResult.err(f"未知操作：{action}，支持的操作为：list, get_details, execute, execute_script")
        
        except Exception as e:
            return ToolResult.err(f"技能管理失败：{e}")

    def _execute_skill(self, skill_config, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """执行技能"""
        try:
            # 内置技能处理
            if skill_config.type == "builtin" or skill_config.name in [
                "read_file", "write_file", "edit_file", "list_directory",
                "grep", "glob", "calculate", "web_search", "fetch_url",
                "analyze_paper", "open_file", "diff_files", "move_file",
                "copy_file", "file_info", "cli_execute", "config_manage",
                "sub_agent", "hub_search", "hub_lightweight_crawl",
                "hub_discover_subscriptions"
            ]:
                return self._execute_builtin_skill(skill_config.name, parameters)
            
            # 自定义技能处理
            elif skill_config.handler_path:
                return self._execute_custom_skill(skill_config, parameters)
            
            else:
                return {
                    "success": False,
                    "message": f"技能 '{skill_config.name}' 没有可执行的处理器"
                }
        
        except Exception as e:
            return {
                "success": False,
                "message": f"执行技能失败：{e}"
            }

    def _execute_builtin_skill(self, skill_name: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """执行内置技能"""
        return {
            "success": True,
            "data": {"skill": skill_name, "parameters": parameters},
            "message": f"内置技能 '{skill_name}' 应通过对应的工具调用。参数：{json.dumps(parameters, ensure_ascii=False)}"
        }

    def _execute_custom_skill(self, skill_config, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """执行自定义技能"""
        try:
            if skill_config.handler_path:
                module_path, func_name = skill_config.handler_path.rsplit('.', 1)
                module = __import__(module_path, fromlist=[func_name])
                handler = getattr(module, func_name)
                result = handler(parameters)
                
                return {
                    "success": True,
                    "data": result if isinstance(result, dict) else {"result": result},
                    "message": f"技能 '{skill_config.name}' 执行成功"
                }
            else:
                return {
                    "success": False,
                    "message": f"技能 '{skill_config.name}' 没有配置处理器路径"
                }
        except Exception as e:
            return {
                "success": False,
                "message": f"执行自定义技能失败：{e}"
            }

    def register_skill_handler(self, skill_name: str, handler_func: Callable):
        """注册技能处理器"""
        self._skill_registry[skill_name] = handler_func

    def unregister_skill_handler(self, skill_name: str):
        """取消注册技能处理器"""
        if skill_name in self._skill_registry:
            del self._skill_registry[skill_name]

    def list_skill_handlers(self) -> List[str]:
        """列出已注册的技能处理器"""
        return list(self._skill_registry.keys())

    def _execute_skill_script(self, skill_name: str, script_name: str, parameters: Dict[str, Any]) -> ToolResult:
        """执行技能目录下 scripts/ 中的脚本

        根据 SKILL.md 中的 scripts 字段查找脚本路径，
        通过 subprocess 执行，参数以 JSON 字符串传入 stdin。
        如果 SKILL.md 中声明了 api_key_ref，自动注入环境变量。
        """
        import subprocess
        import os

        # 查找技能目录
        skills_dir = self.base_dir / "skills" if (self.base_dir / "skills").exists() else Path(__file__).parent.parent / "skills"
        skill_dir = skills_dir / skill_name
        if not skill_dir.exists():
            return ToolResult.err(f"技能目录不存在: {skill_dir}")

        # 解析 SKILL.md 获取 scripts 和 api_key_ref
        script_path = None
        api_key_env = ""
        try:
            from .skill_system import Skill
            skill_obj = Skill.from_skill_md(skill_dir)
            if skill_obj:
                # 从 scripts 字段查找脚本路径
                scripts = skill_obj.scripts or {}
                script_rel = scripts.get(script_name)
                if script_rel:
                    script_path = skill_dir / script_rel
                else:
                    # 回退：查找 scripts/ 目录下的同名脚本
                    script_path = skill_dir / "scripts" / f"{script_name}.py"

                api_key_env = skill_obj.api_key_ref or ""
        except Exception:
            # 回退：查找 scripts/ 目录
            script_path = skill_dir / "scripts" / f"{script_name}.py"

        if not script_path or not script_path.exists():
            return ToolResult.err(f"脚本不存在: {script_path} (skill={skill_name}, script={script_name})")

        # 构建环境变量
        env = dict(os.environ)
        # api_key_ref 指定的环境变量已经存在于 os.environ 中，自动继承
        # 如果未设置，给出提示
        if api_key_env and not os.environ.get(api_key_env):
            return ToolResult.err(
                f"技能 '{skill_name}' 需要环境变量 '{api_key_env}'，"
                f"请先设置后再执行。"
            )

        # 执行脚本
        try:
            params_json = json.dumps(parameters, ensure_ascii=False)
            proc = subprocess.run(
                [sys.executable, str(script_path), params_json],
                capture_output=True,
                text=True,
                timeout=60,
                env=env,
                cwd=str(skill_dir),
            )

            if proc.returncode != 0:
                error_msg = proc.stderr.strip() or proc.stdout.strip()
                return ToolResult.err(f"脚本执行失败 (exit={proc.returncode}): {error_msg}")

            output = proc.stdout.strip()
            return ToolResult.ok(
                data={"skill": skill_name, "script": script_name, "output": output},
                message=output or f"脚本 '{script_name}' 执行成功（无输出）"
            )

        except subprocess.TimeoutExpired:
            return ToolResult.err(f"脚本执行超时 (60s)")
        except Exception as e:
            return ToolResult.err(f"脚本执行异常: {e}")


class RAGTool(Tool):
    """RAG（检索增强生成）工具 - 文档向量化存储和检索"""
    name = "rag_retrieval"
    category = "search_code"
    keywords = ["rag", "retrieval", "检索", "知识库", "向量"]
    model_hint = "向知识库添加文档或检索相关内容时使用。"
    description = "管理向量知识库：添加文档、检索相关内容。"
    parameters = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "description": "操作类型：add_file=添加文件, add_directory=添加目录, retrieve=检索, query=查询, get_count=计数, list_documents=列文档, clear=清空",
                "enum": ["add_file", "add_directory", "retrieve", "query", "get_count", "list_documents", "clear"],
                "example": "retrieve"
            },
            "file_path": {"type": "string", "description": "[add_file] 文件路径", "example": "docs/readme.md"},
            "directory_path": {"type": "string", "description": "[add_directory] 目录路径", "example": "docs/"},
            "query": {"type": "string", "description": "[retrieve/query] 查询文本", "example": "如何配置系统"},
            "top_k": {"type": "integer", "description": "返回结果数量", "default": 4, "example": 5},
            "metadata": {"type": "object", "description": "[add_file] 元数据"},
        },
        "required": ["action"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()
        self._rag_engine = None

    def _get_rag_engine(self):
        """获取或初始化 RAG 引擎"""
        if self._rag_engine is None:
            try:
                from .rag.rag_engine import RAGEngine
                persist_dir = self.base_dir / "rag_data"
                persist_dir.mkdir(exist_ok=True)
                self._rag_engine = RAGEngine(persist_directory=persist_dir)
            except Exception as e:
                logger.warning(f"RAG引擎初始化失败: {e}")
                return None
        return self._rag_engine

    def execute(self, action: str, file_path: Optional[str] = None, directory_path: Optional[str] = None,
                query: Optional[str] = None, top_k: int = 4, metadata: Optional[Dict[str, Any]] = None) -> str:
        result = self.execute_structured(action, file_path, directory_path, query, top_k, metadata)
        return result.to_json()

    def execute_structured(self, action: str, file_path: Optional[str] = None, directory_path: Optional[str] = None,
                          query: Optional[str] = None, top_k: int = 4, metadata: Optional[Dict[str, Any]] = None) -> ToolResult:
        try:
            rag_engine = self._get_rag_engine()
            if rag_engine is None:
                return ToolResult.err("RAG引擎初始化失败")

            if action == "add_file":
                if not file_path:
                    return ToolResult.err("add_file 需要 file_path 参数")
                
                file_path_obj = Path(file_path)
                if not file_path_obj.is_absolute():
                    file_path_obj = self.base_dir / file_path_obj
                
                if not file_path_obj.exists():
                    return ToolResult.err(f"文件不存在: {file_path}")
                
                chunk_ids = rag_engine.add_file(file_path_obj, metadata=metadata)
                return ToolResult.ok(
                    data={"file": str(file_path_obj), "chunk_ids": chunk_ids, "chunk_count": len(chunk_ids)},
                    message=f"✅ 已添加文件: {file_path_obj.name}，生成 {len(chunk_ids)} 个文本块"
                )
            
            elif action == "add_directory":
                if not directory_path:
                    return ToolResult.err("add_directory 需要 directory_path 参数")
                
                dir_path = Path(directory_path)
                if not dir_path.is_absolute():
                    dir_path = self.base_dir / dir_path
                
                if not dir_path.exists() or not dir_path.is_dir():
                    return ToolResult.err(f"目录不存在: {directory_path}")
                
                result = rag_engine.add_directory(dir_path)
                return ToolResult.ok(
                    data={"directory": str(dir_path), "result": result},
                    message=f"✅ 已添加目录: {directory_path}，共 {result['file_count']} 个文件，{result['chunk_count']} 个文本块"
                )
            
            elif action == "retrieve":
                if not query:
                    return ToolResult.err("retrieve 需要 query 参数")
                
                results = rag_engine.retrieve_with_scores(query, top_k=top_k)
                lines = [f"=== 检索到 {len(results)} 个相关片段 ==="]
                for i, (doc, score) in enumerate(results, 1):
                    content_preview = doc.page_content[:200].replace('\n', ' ')
                    lines.append(f"\n【片段 {i}】(相似度: {score:.3f})")
                    lines.append(f"内容: {content_preview}...")
                    if doc.metadata:
                        lines.append(f"元数据: {json.dumps(doc.metadata, ensure_ascii=False)}")
                
                return ToolResult.ok(
                    data={"query": query, "results": [(doc.page_content, score) for doc, score in results]},
                    message="\n".join(lines)
                )
            
            elif action == "query":
                if not query:
                    return ToolResult.err("query 需要 query 参数")
                
                result = rag_engine.query(query, top_k=top_k)
                lines = [f"=== 查询结果 ({result['count']} 个文档) ==="]
                for doc_info in result["documents"]:
                    lines.append(f"\n【来源: {doc_info['source']}】")
                    lines.append(f"内容: {doc_info['content'][:300]}...")
                
                return ToolResult.ok(
                    data=result,
                    message="\n".join(lines)
                )
            
            elif action == "get_count":
                count = rag_engine.get_document_count()
                return ToolResult.ok(
                    data={"document_count": count},
                    message=f"📊 向量库中共有 {count} 个文档"
                )
            
            elif action == "list_documents":
                docs = rag_engine.get_all_documents()
                lines = [f"=== 所有文档 ({len(docs)} 个) ==="]
                for i, doc in enumerate(docs, 1):
                    lines.append(f"\n【{i}】{doc.source}")
                    lines.append(f"内容预览: {doc.page_content[:100]}...")
                
                return ToolResult.ok(
                    data={"documents": [{"source": doc.source, "id": doc.id} for doc in docs]},
                    message="\n".join(lines)
                )
            
            elif action == "clear":
                rag_engine.clear()
                return ToolResult.ok(
                    data={"status": "cleared"},
                    message="✅ 已清空向量库中的所有文档"
                )
            
            else:
                return ToolResult.err(f"未知操作：{action}")
        
        except Exception as e:
            return ToolResult.err(f"RAG操作失败: {e}")


class SandboxTool(Tool):
    """沙箱执行工具 - 安全执行命令"""
    name = "sandbox_execute"
    category = "system"
    keywords = ["sandbox", "execute", "沙盒", "安全执行"]
    model_hint = "在受限安全环境中执行可能不信任的命令时使用。"
    description = "在受限沙箱环境中安全执行命令。"
    parameters = {
        "type": "object",
        "properties": {
            "command": {"type": "string", "description": "要执行的命令", "example": "python -c 'print(1+1)'"},
            "allow_network": {"type": "boolean", "description": "允许网络访问", "default": False},
            "max_time": {"type": "integer", "description": "最大执行时间（秒）", "default": 30, "example": 10},
            "open_terminal": {"type": "boolean", "description": "在新终端窗口执行", "default": False},
            "keep_terminal_open": {"type": "boolean", "description": "执行后保持终端打开", "default": True},
        },
        "required": ["command"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, command: str, allow_network: bool = False, max_time: int = 30, open_terminal: bool = False, keep_terminal_open: bool = True) -> str:
        result = self.execute_structured(command, allow_network, max_time, open_terminal, keep_terminal_open)
        return result.to_json()

    def execute_structured(self, command: str, allow_network: bool = False, max_time: int = 30, open_terminal: bool = False, keep_terminal_open: bool = True) -> ToolResult:
        try:
            if open_terminal:
                return self._execute_in_terminal(command, keep_terminal_open)
            
            try:
                from .sandbox.executor import SandboxExecutor
                from .sandbox.policy import SandboxPolicy
                
                policy = SandboxPolicy(
                    allow_network=allow_network,
                    max_execution_time=max_time
                )
                executor = SandboxExecutor(policy=policy, cwd=str(self.base_dir))
                
                result = executor.execute(command)
                
                output_lines = [
                    f"$ {command}",
                    result.stdout or "(无输出)",
                ]
                if result.stderr:
                    output_lines.append(f"错误: {result.stderr}")
                output_lines.append(f"[返回码: {result.exit_code}]")
                if result.timed_out:
                    output_lines.append("[超时]")
                output_lines.append(f"[耗时: {result.duration_ms}ms]")
                
                if result.success:
                    return ToolResult.ok(
                        data={
                            "command": command,
                            "exit_code": result.exit_code,
                            "stdout": result.stdout,
                            "stderr": result.stderr,
                            "duration_ms": result.duration_ms,
                        },
                        message="\n".join(output_lines)
                    )
                else:
                    return ToolResult.ok(
                        data={
                            "command": command,
                            "exit_code": result.exit_code,
                            "stdout": result.stdout,
                            "stderr": result.stderr,
                            "duration_ms": result.duration_ms,
                            "error": result.error,
                            "timed_out": result.timed_out,
                        },
                        message="\n".join(output_lines)
                    )
            except ImportError:
                # 如果没有沙箱模块，退回到普通的CLI执行
                logger.warning("沙箱模块未找到，使用普通CLI执行")
                from .tools import CLITool
                cli_tool = CLITool(base_dir=self.base_dir)
                return cli_tool.execute_structured(command=command)
        
        except Exception as e:
            return ToolResult.err(f"沙箱执行失败: {e}")

    def _execute_in_terminal(self, command: str, keep_open: bool) -> ToolResult:
        import subprocess
        import platform

        working_dir = self.base_dir
        working_dir.mkdir(parents=True, exist_ok=True)

        try:
            system = platform.system()
            
            if system == "Windows":
                cmd_parts = ["cmd", "/k"] if keep_open else ["cmd", "/c"]
                cmd_parts.append(command)
                subprocess.Popen(cmd_parts, cwd=str(working_dir), creationflags=subprocess.CREATE_NEW_CONSOLE)
                
            elif system == "Darwin":
                script = f'cd "{working_dir}"; {command}'
                if keep_open:
                    script += '; exec bash'
                applescript = f'tell application "Terminal" to do script "{script}"'
                subprocess.run(["osascript", "-e", applescript], check=True)
                
            else:
                shell = os.environ.get("SHELL", "bash")
                if keep_open:
                    script = f"{command}; exec {shell}"
                else:
                    script = command
                subprocess.Popen([shell, "-c", script], cwd=str(working_dir))

            return ToolResult.ok(
                data={"command": command, "in_terminal": True, "keep_open": keep_open, "sandbox": False},
                message=f"🖥️ 已在新终端中执行命令: {command}\n(注意：终端模式下安全限制可能较宽松)"
            )
        except Exception as e:
            return ToolResult.err(f"在终端执行失败: {e}")


class MCPClientTool(Tool):
    """MCP客户端工具 - 管理外部MCP服务"""
    name = "mcp_client"
    category = "system"
    keywords = ["mcp", "client", "外部服务"]
    model_hint = "列出或调用已连接的外部 MCP 服务工具时使用。"
    description = "管理 MCP 客户端：列出工具、调用工具、查看状态。"
    parameters = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "description": "操作类型：list_tools=列出工具, call_tool=调用工具, get_status=查看状态",
                "enum": ["list_tools", "call_tool", "get_status"],
                "example": "list_tools"
            },
            "tool_name": {"type": "string", "description": "[call_tool] 工具名称", "example": "weather_api"},
            "arguments": {"type": "object", "description": "[call_tool] 工具参数"},
        },
        "required": ["action"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()
        self._client_manager = None

    def _get_client_manager(self):
        """获取或初始化 MCP 客户端管理器"""
        if self._client_manager is None:
            try:
                from .mcp_client.client import MCPClientManager
                self._client_manager = MCPClientManager()
            except Exception as e:
                logger.warning(f"MCP客户端初始化失败: {e}")
                return None
        return self._client_manager

    def execute(self, action: str, tool_name: Optional[str] = None, arguments: Optional[Dict[str, Any]] = None) -> str:
        result = self.execute_structured(action, tool_name, arguments)
        return result.to_json()

    def execute_structured(self, action: str, tool_name: Optional[str] = None, arguments: Optional[Dict[str, Any]] = None) -> ToolResult:
        try:
            client_mgr = self._get_client_manager()
            if client_mgr is None:
                return ToolResult.err("MCP客户端管理器初始化失败")

            if action == "list_tools":
                tools = client_mgr.list_tools()
                lines = ["=== 可用的 MCP 工具 ==="]
                for tool in tools:
                    name = tool.get("name", "未知")
                    desc = tool.get("description", "无描述")
                    lines.append(f"• {name}: {desc}")
                
                return ToolResult.ok(
                    data={"tools": tools, "count": len(tools)},
                    message="\n".join(lines) if lines else "暂无可用的 MCP 工具"
                )
            
            elif action == "call_tool":
                if not tool_name:
                    return ToolResult.err("call_tool 需要 tool_name 参数")
                
                result = client_mgr.call_tool(tool_name, arguments or {})
                return ToolResult.ok(
                    data={"tool": tool_name, "arguments": arguments, "result": result},
                    message=f"✅ MCP 工具 '{tool_name}' 执行成功"
                )
            
            elif action == "get_status":
                clients = client_mgr.list_clients()
                lines = ["=== MCP 客户端状态 ==="]
                for name, info in clients.items():
                    state = info.state.value
                    lines.append(f"• {name}: {state}")
                    if info.error:
                        lines.append(f"  错误: {info.error}")
                
                return ToolResult.ok(
                    data={"clients": {name: {"state": info.state.value, "error": info.error} for name, info in clients.items()}},
                    message="\n".join(lines) if lines else "暂无连接的 MCP 客户端"
                )
            
            else:
                return ToolResult.err(f"未知操作：{action}")
        
        except Exception as e:
            return ToolResult.err(f"MCP客户端操作失败: {e}")


class WorkflowTool(Tool):
    """工作流管理工具 - 定义、执行、管理工作流"""
    name = "workflow"
    category = "workflow"
    keywords = ["workflow", "工作流", "编排", "自动化"]
    model_hint = "定义、启动、暂停、恢复或查看工作流状态时使用。"
    description = "管理工作流：定义、启动、暂停、恢复、取消、查看日志。"
    parameters = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "description": "操作类型：define=定义, start=启动, status=查看状态, pause=暂停, resume=恢复, cancel=取消, list=列出, logs=日志, step_results=步骤结果",
                "enum": ["define", "start", "status", "pause", "resume", "cancel", "list", "logs", "step_results"],
                "example": "list"
            },
            "workflow_id": {"type": "string", "description": "[define/start] 工作流 ID", "example": "data_pipeline"},
            "instance_id": {"type": "string", "description": "[status/pause/resume/cancel/logs/step_results] 实例 ID", "example": "inst_abc123"},
            "definition": {"type": "object", "description": "[define] 工作流定义对象"},
            "input_data": {"type": "object", "description": "[start] 输入数据"},
        },
        "required": ["action"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()
        self._engine = None

    def _get_engine(self):
        if self._engine is None:
            from .workflow_engine import get_workflow_engine
            db_path = self.base_dir / "data" / "workflow.db"
            self._engine = get_workflow_engine(db_path)
        return self._engine

    def execute(self, action: str, workflow_id: Optional[str] = None, 
                instance_id: Optional[str] = None, definition: Optional[Dict] = None,
                input_data: Optional[Dict] = None) -> str:
        result = self.execute_structured(action, workflow_id, instance_id, definition, input_data)
        return result.to_json()

    def execute_structured(self, action: str, workflow_id: Optional[str] = None,
                          instance_id: Optional[str] = None, definition: Optional[Dict] = None,
                          input_data: Optional[Dict] = None) -> ToolResult:
        try:
            from .workflow_engine import WorkflowDefinition, StepDefinition, StepType
            
            engine = self._get_engine()
            
            if action == "define":
                if not workflow_id or not definition:
                    return ToolResult.err("define 需要 workflow_id 和 definition")
                
                steps = [StepDefinition.from_dict(s) for s in definition.get("steps", [])]
                wf_def = WorkflowDefinition(
                    workflow_id=workflow_id,
                    name=definition.get("name", workflow_id),
                    description=definition.get("description", ""),
                    version=definition.get("version", "1.0"),
                    steps=steps,
                    entry_step=definition.get("entry_step"),
                    checkpoint_after=set(definition.get("checkpoint_after", [])),
                )
                engine.persistence.save_definition(wf_def)
                return ToolResult.ok(
                    data={"workflow_id": workflow_id},
                    message=f"✅ 工作流定义保存成功: {wf_def.name}"
                )
                
            elif action == "start":
                if not workflow_id:
                    return ToolResult.err("start 需要 workflow_id")
                
                wf_def = engine.persistence.get_definition(workflow_id)
                if not wf_def:
                    return ToolResult.err(f"工作流不存在: {workflow_id}")
                
                inst_id = engine.start_workflow(wf_def, input_data or {})
                return ToolResult.ok(
                    data={"instance_id": inst_id},
                    message=f"🚀 工作流启动成功: {wf_def.name}\n实例 ID: {inst_id}"
                )
                
            elif action == "status":
                if not instance_id:
                    return ToolResult.err("status 需要 instance_id")
                
                status = engine.get_status(instance_id)
                if not status:
                    return ToolResult.err(f"实例不存在: {instance_id}")
                
                lines = [f"📊 工作流状态: {status['workflow_id']}"]
                lines.append(f"状态: {status['status']}")
                lines.append(f"当前步骤: {status.get('current_step', '-')}")
                lines.append(f"进度: {status.get('progress', 0)*100:.0f}%")
                lines.append(f"启动时间: {status.get('started_at', '-')}")
                
                return ToolResult.ok(
                    data=status,
                    message="\n".join(lines)
                )
                
            elif action == "pause":
                if not instance_id:
                    return ToolResult.err("pause 需要 instance_id")
                
                if engine.pause_workflow(instance_id):
                    return ToolResult.ok(
                        data={"instance_id": instance_id},
                        message="⏸️ 工作流已暂停"
                    )
                else:
                    return ToolResult.err(f"暂停失败，实例可能不存在或未在运行: {instance_id}")
                
            elif action == "resume":
                if not instance_id:
                    return ToolResult.err("resume 需要 instance_id")
                
                if engine.resume_workflow(instance_id):
                    return ToolResult.ok(
                        data={"instance_id": instance_id},
                        message="▶️ 工作流已恢复"
                    )
                else:
                    return ToolResult.err(f"恢复失败，实例可能不存在或不可恢复: {instance_id}")
                
            elif action == "cancel":
                if not instance_id:
                    return ToolResult.err("cancel 需要 instance_id")
                
                if engine.cancel_workflow(instance_id):
                    return ToolResult.ok(
                        data={"instance_id": instance_id},
                        message="🚫 工作流已取消"
                    )
                else:
                    return ToolResult.err(f"取消失败，实例可能不存在或未在运行: {instance_id}")
                
            elif action == "logs":
                if not instance_id:
                    return ToolResult.err("logs 需要 instance_id")
                
                logs = engine.get_logs(instance_id, limit=20)
                if not logs:
                    return ToolResult.ok(message="暂无日志")
                
                lines = [f"📋 工作流日志 ({instance_id[:8]}):"]
                for log in reversed(logs):
                    level = log.get("log_level", "info")
                    ts = (log.get("created_at") or "")[11:19]
                    step = log.get("step_name", "")
                    step_info = f"[{step}] " if step else ""
                    lines.append(f"  {ts} {step_info}{log.get('message', '')}")
                
                return ToolResult.ok(
                    data={"logs": logs},
                    message="\n".join(lines)
                )
                
            elif action == "step_results":
                if not instance_id:
                    return ToolResult.err("step_results 需要 instance_id")
                
                step_results = engine.get_step_results(instance_id)
                if not step_results:
                    return ToolResult.ok(message="暂无步骤结果")
                
                lines = [f"📊 步骤结果 ({instance_id[:8]}):"]
                for sid, sr in step_results.items():
                    status = sr.get("status", "")
                    icon = {"completed": "✅", "running": "🔄", "failed": "❌"}.get(status, "❓")
                    lines.append(f"  {icon} {sid}: {status}")
                    if sr.get("error"):
                        lines.append(f"     错误: {sr['error']}")
                
                return ToolResult.ok(
                    data={"step_results": step_results},
                    message="\n".join(lines)
                )
                
            elif action == "list":
                defs = engine.persistence.list_definitions()
                instances = engine.persistence.list_instances()
                
                lines = ["📋 工作流定义:"]
                if defs:
                    for d in defs:
                        lines.append(f"  • {d['workflow_id']} - {d['name']}")
                else:
                    lines.append("  (暂无定义)")
                
                lines.append("\n🔄 最近运行实例:")
                if instances:
                    for i in instances[:10]:
                        lines.append(f"  • {i['instance_id']} - {i['workflow_id']} - {i['status']}")
                else:
                    lines.append("  (暂无运行)")
                
                return ToolResult.ok(
                    data={"definitions": defs, "instances": instances},
                    message="\n".join(lines)
                )
                
            else:
                return ToolResult.err(f"未知操作: {action}")
        
        except Exception as e:
            import traceback
            logger.error(f"Workflow tool error: {e}")
            logger.error(traceback.format_exc())
            return ToolResult.err(f"工作流操作失败: {e}")


class ApprovalTool(Tool):
    """审批管理工具 - 查看和处理待审批的请求"""
    name = "approval"
    category = "system"
    keywords = ["approval", "审批", "确认"]
    model_hint = "查看待审批请求或批准/拒绝操作时使用。"
    description = "审批管理：查看待审批、批准/拒绝请求、设置审批模式。"
    parameters = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "description": "操作类型：list_pending=待审批, approve=批准, deny=拒绝, always_approve=始终批准, set_mode=设置模式",
                "enum": ["list_pending", "approve", "deny", "always_approve", "set_mode"],
                "example": "list_pending"
            },
            "request_id": {"type": "string", "description": "[approve/deny/always_approve] 请求 ID", "example": "req_abc123"},
            "mode": {"type": "string", "description": "[set_mode] 审批模式", "enum": ["suggest", "auto_edit", "full_auto"], "example": "auto_edit"},
        },
        "required": ["action"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()
        self._manager = None

    def _get_manager(self):
        if self._manager is None:
            from .harness import ApprovalManager, ApprovalMode
            self._manager = ApprovalManager(mode=ApprovalMode.SUGGEST)
        return self._manager

    def execute(self, action: str, request_id: Optional[str] = None, mode: Optional[str] = None) -> str:
        result = self.execute_structured(action, request_id, mode)
        return result.to_json()

    def execute_structured(self, action: str, request_id: Optional[str] = None,
                          mode: Optional[str] = None) -> ToolResult:
        try:
            from .harness import ApprovalManager, ApprovalMode, ApprovalDecision
            
            mgr = self._get_manager()
            
            if action == "list_pending":
                pending = mgr.get_pending()
                
                lines = ["⏳ 待审批请求:"]
                if pending:
                    for req in pending:
                        lines.append(f"  [{req.id}] {req.tool_name}")
                        lines.append(f"    风险等级: {req.risk_level}")
                        if req.reason:
                            lines.append(f"    原因: {req.reason}")
                else:
                    lines.append("  (暂无待审批)")
                
                return ToolResult.ok(
                    data={"pending": [{"id": r.id, "tool": r.tool_name, "risk": r.risk_level} for r in pending]},
                    message="\n".join(lines)
                )
                
            elif action in ("approve", "deny", "always_approve"):
                if not request_id:
                    return ToolResult.err(f"{action} 需要 request_id")
                
                decision_map = {
                    "approve": ApprovalDecision.APPROVE,
                    "deny": ApprovalDecision.DENY,
                    "always_approve": ApprovalDecision.ALWAYS_APPROVE,
                }
                
                mgr.decide(request_id, decision_map[action])
                
                msg = {
                    "approve": "✅ 请求已批准",
                    "deny": "❌ 请求已拒绝",
                    "always_approve": "✅ 请求已批准，此类请求将自动通过",
                }[action]
                
                return ToolResult.ok(
                    data={"request_id": request_id, "action": action},
                    message=msg
                )
                
            elif action == "set_mode":
                if not mode:
                    return ToolResult.err("set_mode 需要 mode 参数")
                
                mode_map = {
                    "suggest": ApprovalMode.SUGGEST,
                    "auto_edit": ApprovalMode.AUTO_EDIT,
                    "full_auto": ApprovalMode.FULL_AUTO,
                }
                mgr._mode = mode_map[mode]
                
                mode_desc = {
                    "suggest": "建议模式 - 所有操作需要审批",
                    "auto_edit": "自动编辑模式 - 仅敏感操作需要审批",
                    "full_auto": "全自动模式 - 无需审批",
                }
                
                return ToolResult.ok(
                    data={"mode": mode},
                    message=f"🔧 审批模式已设置: {mode_desc[mode]}"
                )
                
            else:
                return ToolResult.err(f"未知操作: {action}")
        
        except Exception as e:
            return ToolResult.err(f"审批操作失败: {e}")


class ExtensionSkillTool(Tool):
    """扩展技能管理工具 - 基于新的 MCP 扩展系统"""
    name = "extension_skill"
    category = "system"
    keywords = ["extension", "skill", "扩展技能"]
    model_hint = "管理 MCP 扩展系统中的技能：列出、获取、执行、加载、导出。"
    description = "管理扩展技能：列出、获取、执行、加载、导出。"
    parameters = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "description": "【必填】操作类型：list=列出技能, get=获取详情, execute=执行技能, load=从文件加载技能, export=导出技能到文件",
                "enum": ["list", "get", "execute", "load", "export"],
                "example": "list"
            },
            "skill_name": {"type": "string", "description": "【get/execute/export 必填】技能名称", "example": "my_skill"},
            "file_path": {"type": "string", "description": "【load/export 必填】技能文件的路径", "example": "skills/my_skill.json"},
            "parameters": {"type": "object", "description": "【execute 必填】传递给技能的执行参数字典"},
            "category": {"type": "string", "description": "【list 可选】按分类过滤技能列表", "example": "web"},
        },
        "required": ["action"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()
        self._registry = None

    def _get_registry(self):
        if self._registry is None:
            from .extensions.skills import get_skill_registry
            self._registry = get_skill_registry()
        return self._registry

    def execute(self, action: str, skill_name: Optional[str] = None,
                file_path: Optional[str] = None,
                parameters: Optional[Dict] = None,
                category: Optional[str] = None) -> str:
        result = self.execute_structured(action, skill_name, file_path, parameters, category)
        return result.to_json()

    def execute_structured(self, action: str, skill_name: Optional[str] = None,
                          file_path: Optional[str] = None,
                          parameters: Optional[Dict] = None,
                          category: Optional[str] = None) -> ToolResult:
        try:
            registry = self._get_registry()
            
            if action == "list":
                skills = registry.list_skills(category)
                lines = ["📋 技能列表:"]
                for skill in skills:
                    status = "✅" if skill.enabled else "❌"
                    lines.append(f"  {status} {skill.name} - {skill.description} ({skill.category}")
                return ToolResult.ok(
                    data={"skills": [{"name": s.name, "description": s.description, "category": s.category, "enabled": s.enabled} for s in skills]},
                    message="\n".join(lines)
                )
            
            elif action == "get":
                if not skill_name:
                    return ToolResult.err("get 需要 skill_name")
                skill = registry.get_skill(skill_name)
                if not skill:
                    return ToolResult.err(f"技能 '{skill_name}' 不存在")
                return ToolResult.ok(
                    data={"name": skill.name, "description": skill.description, "parameters": [p.__dict__ for p in skill.parameters]},
                    message=f"技能 '{skill.name}: {skill.description}"
                )
            
            elif action == "execute":
                if not skill_name:
                    return ToolResult.err("execute 需要 skill_name")
                import asyncio
                result = asyncio.run(registry.execute_skill(skill_name, **(parameters or {})))
                return ToolResult.ok(
                    data={"result": result},
                    message=f"技能执行结果: {result}"
                )
            
            elif action == "load":
                if not file_path:
                    return ToolResult.err("load 需要 file_path")
                path = Path(file_path)
                if not path.is_file():
                    return ToolResult.ok(message="成功加载技能文件")
                else:
                    return ToolResult.err("文件不存在")
            
            elif action == "export":
                if not skill_name or not file_path:
                    return ToolResult.err("export 需要 skill_name 和 file_path")
                path = Path(file_path)
                registry.export_skill(skill_name, path)
                return ToolResult.ok(message=f"已导出技能")
            
            else:
                return ToolResult.err(f"未知操作")
        
        except Exception as e:
            import traceback
            import logging
            logging.getLogger(__name__).error(traceback.format_exc())
            return ToolResult.err(f"技能操作失败: {e}")


class SessionTool(Tool):
    """会话管理工具 - 管理多个对话会话"""
    name = "session"
    category = "system"
    keywords = ["session", "会话", "对话管理"]
    model_hint = "创建、切换、查看或删除对话会话时使用。"
    description = "管理多个对话会话：创建、切换、查看、删除。"
    parameters = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "description": "操作类型：list=列出, create=创建, set_active=激活, get=获取, add_message=添加消息, delete=删除",
                "enum": ["list", "create", "set_active", "get", "add_message", "delete"],
                "example": "list"
            },
            "name": {"type": "string", "description": "[create] 会话名称", "example": "代码分析会话"},
            "session_id": {"type": "string", "description": "[set_active/get/add_message/delete] 会话 ID", "example": "sess_abc123"},
            "role": {"type": "string", "description": "[add_message] 消息角色", "enum": ["user", "assistant", "system"], "example": "user"},
            "content": {"type": "string", "description": "[add_message] 消息内容", "example": "请分析这段代码"},
        },
        "required": ["action"],
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()
        self._manager = None

    def _get_manager(self):
        if self._manager is None:
            from .extensions.unified_session import get_unified_session_manager
            self._manager = get_unified_session_manager()
        return self._manager

    def execute(self, action: str, name: Optional[str] = None,
                session_id: Optional[str] = None,
                role: Optional[str] = None,
                content: Optional[str] = None) -> str:
        result = self.execute_structured(action, name, session_id, role, content)
        return result.to_json()

    def execute_structured(self, action: str, name: Optional[str] = None,
                        session_id: Optional[str] = None,
                        role: Optional[str] = None,
                        content: Optional[str] = None) -> ToolResult:
        try:
            manager = self._get_manager()
            
            if action == "list":
                sessions = manager.list_sessions()
                active_id = manager.get_active_session_id()
                lines = ["💬 会话列表:"]
                for s in sessions:
                    marker = "👉" if s.session_id == active_id else "  "
                    task_info = " 🔄" if s.has_running_tasks else ""
                    lines.append(f"  {marker} {s.title} ({s.message_count} 条消息, {s.turn_count} 轮){task_info}")
                return ToolResult.ok(
                    data={"sessions": [s.to_dict() for s in sessions], "active_session_id": active_id},
                    message="\n".join(lines)
                )
            
            elif action == "create":
                if not name:
                    return ToolResult.err("create 需要 name")
                sid = manager.create_session(name)
                if sid:
                    return ToolResult.ok(
                        data={"session_id": sid, "name": name},
                        message=f"已创建会话: {name}"
                    )
                else:
                    return ToolResult.err("创建会话失败")
            
            elif action == "set_active":
                if not session_id:
                    return ToolResult.err("set_active 需要 session_id")
                if manager.set_active_session(session_id):
                    session = manager.get_session(session_id)
                    title = session.get("title", session_id) if session else session_id
                    return ToolResult.ok(
                        data={"session_id": session_id},
                        message=f"已切换到会话: {title}"
                    )
                else:
                    return ToolResult.err("会话不存在")
            
            elif action == "get":
                if session_id:
                    session = manager.get_session(session_id)
                else:
                    session = manager.get_active_session()
                if not session:
                    return ToolResult.err("没有活动会话")
                title = session.get("title", "未命名")
                msg_count = len(session.get("messages", []))
                return ToolResult.ok(
                    data=session,
                    message=f"当前会话: {title} ({msg_count} 条消息)"
                )
            
            elif action == "add_message":
                if not role or not content:
                    return ToolResult.err("add_message 需要 role 和 content")
                if not session_id:
                    active_id = manager.get_active_session_id()
                    if not active_id:
                        return ToolResult.err("没有活动会话，请指定 session_id")
                    session_id = active_id
                if manager.add_message(session_id, role, content):
                    return ToolResult.ok(message="已添加消息")
                else:
                    return ToolResult.err("添加消息失败")
            
            elif action == "delete":
                if not session_id:
                    return ToolResult.err("delete 需要 session_id")
                if manager.delete_session(session_id):
                    return ToolResult.ok(message="已删除会话")
                else:
                    return ToolResult.err("会话不存在")
            
            else:
                return ToolResult.err(f"未知操作: {action}")
        
        except Exception as e:
            import traceback
            import logging
            logging.getLogger(__name__).error(traceback.format_exc())
            return ToolResult.err(f"会话操作失败: {e}")


# ============================================================
# SearchTool 搜索引擎系统 - 替代全量数据返回
# ============================================================

class SearchTool(Tool):
    """搜索工具基类：支持关键词搜索、标签过滤、分页、摘要模式"""

    default_limit: int = 10
    item_type_label: str = "结果"
    type_label: str = "类型"

    def _format_header(self, total: int, keyword: str, offset: int, limit: int) -> str:
        """生成摘要头行"""
        parts = []
        if total == 0:
            return f"🔍 未找到匹配的{self.item_type_label}"
        parts.append(f"🔍 找到 {total} 个{self.item_type_label}")
        if keyword:
            parts.append(f"关键词 '{keyword}'")
        end = offset + limit
        if end < total:
            parts.append(f"显示前 {min(limit, total)} 个")
        else:
            parts.append(f"显示 {offset + 1}-{total} 个")
        return "，".join(parts)

    def _format_footer(self, offset: int, limit: int, total: int) -> Optional[str]:
        """生成分页提示"""
        if offset + limit < total:
            return f"📄 使用 offset={offset + limit} 查看下一页"
        return None

    def _match_keyword(self, text: str, keyword: str) -> bool:
        """关键词模糊匹配（忽略大小写）"""
        if not keyword:
            return True
        return keyword.lower() in text.lower()

    def _match_tags(self, item_tags: Union[List[str], str, None], tag_filter: Optional[str]) -> bool:
        """标签/类别过滤"""
        if not tag_filter:
            return True
        if not item_tags:
            return False
        if isinstance(item_tags, str):
            item_tags = [item_tags]
        tag_filter_lower = tag_filter.lower()
        return any(tag_filter_lower in t.lower() for t in item_tags)

    def _build_result(self, items: list, total: int, keyword: str, offset: int, limit: int) -> ToolResult:
        """构建标准搜索结果"""
        lines = [self._format_header(total, keyword, offset, limit)]
        for item in items:
            lines.append(self._format_item(item))
        footer = self._format_footer(offset, limit, total)
        if footer:
            lines.append(footer)
        message = "\n".join(lines)
        return ToolResult.ok(
            data={
                "total": total,
                "offset": offset,
                "limit": limit,
                "returned": len(items),
                "keyword": keyword,
                "items": items,
            },
            message=message,
        )

    def _format_item(self, item: dict) -> str:
        """子类必须实现：格式化单个条目为摘要行"""
        raise NotImplementedError


# ---------- SearchFilesTool ----------

class SearchFilesTool(SearchTool):
    """文件搜索工具：替代 list_directory 的全量返回，支持文件名/扩展名/路径搜索"""
    name = "search_files"
    category = "file_search"
    keywords = ["search", "files", "文件搜索", "查找文件"]
    model_hint = "当需要搜索文件（按名称、扩展名、路径关键词）而非浏览目录时使用。"
    description = "搜索工作目录中的文件，支持按名称、扩展名、路径关键词过滤，返回摘要列表。"
    parameters = {
        "type": "object",
        "properties": {
            "keyword": {"type": "string", "description": "搜索关键词（匹配文件名、路径）", "example": "test"},
            "extension": {"type": "string", "description": "按扩展名过滤（如 .py, .md）", "example": ".py"},
            "path_contains": {"type": "string", "description": "路径包含关键词", "example": "src/"},
            "min_size": {"type": "integer", "description": "最小文件大小（字节）", "example": 100},
            "limit": {"type": "integer", "description": "每页返回数量，默认 10", "default": 10},
            "offset": {"type": "integer", "description": "分页偏移量，默认 0", "default": 0},
            "recursive": {"type": "boolean", "description": "是否递归搜索子目录", "default": True},
        },
        "required": [],
    }
    item_type_label = "文件"

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, keyword: Optional[str] = None, extension: Optional[str] = None,
                path_contains: Optional[str] = None, min_size: Optional[int] = None,
                limit: int = 10, offset: int = 0, recursive: bool = True) -> str:
        result = self.execute_structured(keyword, extension, path_contains, min_size, limit, offset, recursive)
        return result.to_json()

    def execute_structured(self, keyword: Optional[str] = None, extension: Optional[str] = None,
                           path_contains: Optional[str] = None, min_size: Optional[int] = None,
                           limit: int = 10, offset: int = 0, recursive: bool = True) -> ToolResult:
        try:
            import fnmatch

            def _collect_files(current: Path):
                matched = []
                try:
                    items = sorted(current.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
                except PermissionError:
                    return matched
                for item in items:
                    if item.name.startswith(".") and current != self.base_dir:
                        continue
                    if item.is_dir() and recursive:
                        matched.extend(_collect_files(item))
                    elif item.is_file():
                        # 扩展名过滤
                        if extension and not item.name.lower().endswith(extension.lower()):
                            continue
                        # 路径过滤
                        rel = str(item.relative_to(self.base_dir))
                        if path_contains and path_contains.lower() not in rel.lower():
                            continue
                        # 关键词过滤
                        if not self._match_keyword(item.name, keyword or "") and not self._match_keyword(rel, keyword or ""):
                            continue
                        # 大小过滤
                        if min_size is not None and item.stat().st_size < min_size:
                            continue
                        matched.append(item)
                return matched

            all_files = _collect_files(self.base_dir)
            total = len(all_files)

            # 分页
            page = all_files[offset:offset + limit]
            summary_items = []
            for f in page:
                try:
                    stat = f.stat()
                    size_str = self._format_size(stat.st_size)
                    from datetime import datetime as _dt
                    mtime = _dt.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M')
                except Exception:
                    size_str = "?"
                    mtime = "?"
                summary_items.append({
                    "name": f.name,
                    "path": str(f.relative_to(self.base_dir)),
                    "size": size_str,
                    "size_bytes": stat.st_size if stat.st_size else 0,
                    "modified": mtime,
                    "type": "file",
                    "extension": f.suffix or "(无扩展名)",
                })

            return self._build_result(summary_items, total, keyword or "", offset, limit)

        except Exception as e:
            return ToolResult.err(f"文件搜索失败：{e}")

    def _format_item(self, item: dict) -> str:
        return f"  📄 {item['name']}  [{item['extension']}, {item['size']}]  {item['path']}"

    def _format_size(self, size: int) -> str:
        if size < 1024:
            return f"{size}B"
        elif size < 1024 * 1024:
            return f"{size / 1024:.1f}KB"
        elif size < 1024 * 1024 * 1024:
            return f"{size / (1024 * 1024):.1f}MB"
        else:
            return f"{size / (1024 * 1024 * 1024):.1f}GB"


# ---------- SearchCoursesTool ----------

class SearchCoursesTool(SearchTool):
    """课程搜索工具：替代 ws2_list，支持课程名称搜索"""
    name = "search_courses"
    category = "course_search"
    keywords = ["search", "courses", "课程搜索", "查找课程", "ws2"]
    model_hint = "当需要搜索课程（按名称关键词）而非列出全部课程时使用。"
    description = "搜索 WS2 课程，支持按名称关键词过滤，返回摘要列表。"
    parameters = {
        "type": "object",
        "properties": {
            "keyword": {"type": "string", "description": "课程名称关键词", "example": "机器学习"},
            "limit": {"type": "integer", "description": "每页返回数量，默认 10", "default": 10},
            "offset": {"type": "integer", "description": "分页偏移量，默认 0", "default": 0},
        },
        "required": [],
    }
    item_type_label = "课程"

    def execute(self, keyword: Optional[str] = None, limit: int = 10, offset: int = 0) -> str:
        result = self.execute_structured(keyword, limit, offset)
        return result.to_json()

    def execute_structured(self, keyword: Optional[str] = None, limit: int = 10, offset: int = 0) -> ToolResult:
        try:
            from ws2_data_hub import get_data_hub
            hub = get_data_hub()
            if hub is None:
                return ToolResult.err("数据枢纽未初始化")

            all_courses = hub.list_courses()
            if keyword:
                matched = [c for c in all_courses if self._match_keyword(c.get('title', '') + ' ' + c.get('description', ''), keyword)]
            else:
                matched = all_courses

            total = len(matched)
            page = matched[offset:offset + limit]

            summary_items = []
            for c in page:
                summary_items.append({
                    "title": c.get('title', '未知'),
                    "description": (c.get('description', '') or '')[:80],
                    "lesson_count": c.get('lesson_count', 0),
                    "type": "course",
                    "id": c.get('id', ''),
                })

            return self._build_result(summary_items, total, keyword or "", offset, limit)

        except ImportError:
            return ToolResult.err("ws2_data_hub 模块不可用")
        except Exception as e:
            return ToolResult.err(f"课程搜索失败：{e}")

    def _format_item(self, item: dict) -> str:
        return f"  📚 {item['title']}  [{item['lesson_count']}课时]  {item['description']}"


# ---------- SearchSessionsTool ----------

class SearchSessionsTool(SearchTool):
    """会话搜索工具：替代 session list，支持会话标题/内容搜索"""
    name = "search_sessions"
    category = "session_search"
    keywords = ["search", "sessions", "会话搜索", "查找会话"]
    model_hint = "当需要搜索历史会话（按标题关键词）而非列出全部会话时使用。"
    description = "搜索对话会话，支持按标题关键词过滤，返回摘要列表。"
    parameters = {
        "type": "object",
        "properties": {
            "keyword": {"type": "string", "description": "会话标题关键词", "example": "代码分析"},
            "limit": {"type": "integer", "description": "每页返回数量，默认 10", "default": 10},
            "offset": {"type": "integer", "description": "分页偏移量，默认 0", "default": 0},
        },
        "required": [],
    }
    item_type_label = "会话"

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()
        self._manager = None

    def _get_manager(self):
        if self._manager is None:
            from .extensions.unified_session import get_unified_session_manager
            self._manager = get_unified_session_manager()
        return self._manager

    def execute(self, keyword: Optional[str] = None, limit: int = 10, offset: int = 0) -> str:
        result = self.execute_structured(keyword, limit, offset)
        return result.to_json()

    def execute_structured(self, keyword: Optional[str] = None, limit: int = 10, offset: int = 0) -> ToolResult:
        try:
            manager = self._get_manager()
            all_sessions = manager.list_sessions()

            if keyword:
                matched = [s for s in all_sessions if self._match_keyword(s.title, keyword)]
            else:
                matched = all_sessions

            total = len(matched)
            page = matched[offset:offset + limit]
            active_id = manager.get_active_session_id()

            summary_items = []
            for s in page:
                is_active = s.session_id == active_id
                summary_items.append({
                    "title": s.title,
                    "message_count": s.message_count,
                    "turn_count": s.turn_count,
                    "session_id": s.session_id,
                    "is_active": is_active,
                    "has_running_tasks": s.has_running_tasks,
                    "type": "session",
                    "active_marker": "👉 当前" if is_active else "",
                })

            return self._build_result(summary_items, total, keyword or "", offset, limit)

        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"会话搜索失败: {e}")
            return ToolResult.err(f"会话搜索失败：{e}")

    def _format_item(self, item: dict) -> str:
        marker = item.get('active_marker', '')
        marker = f" {marker}" if marker else ""
        task = " 🔄" if item.get('has_running_tasks') else ""
        return f"  💬 {item['title']}  [{item['message_count']}条, {item['turn_count']}轮]{marker}{task}  {item['session_id']}"


# ---------- SearchConfigsTool ----------

class SearchConfigsTool(SearchTool):
    """配置搜索工具：替代 config get_all，支持配置项名称/值搜索"""
    name = "search_configs"
    category = "config_search"
    keywords = ["search", "configs", "配置搜索", "查找配置"]
    model_hint = "当需要搜索配置项（按名称或值关键词）而非列出全部配置时使用。"
    description = "搜索系统配置项，支持按名称、值关键词过滤，返回摘要列表。"
    parameters = {
        "type": "object",
        "properties": {
            "keyword": {"type": "string", "description": "配置项名称或值的关键词", "example": "theme"},
            "category": {"type": "string", "description": "按类别过滤（可选）", "example": "theme"},
            "limit": {"type": "integer", "description": "每页返回数量，默认 10", "default": 10},
            "offset": {"type": "integer", "description": "分页偏移量，默认 0", "default": 0},
        },
        "required": [],
    }
    item_type_label = "配置项"

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, keyword: Optional[str] = None, category: Optional[str] = None,
                limit: int = 10, offset: int = 0) -> str:
        result = self.execute_structured(keyword, category, limit, offset)
        return result.to_json()

    def execute_structured(self, keyword: Optional[str] = None, category: Optional[str] = None,
                           limit: int = 10, offset: int = 0) -> ToolResult:
        try:
            from .config import get_config_manager
            config_mgr = get_config_manager()
            settings = config_mgr.settings

            # 转为列表
            config_list = [{"key": k, "value": v, "category": self._guess_category(k)} for k, v in settings.items()]

            # 过滤
            matched = config_list
            if keyword:
                matched = [c for c in matched if self._match_keyword(c['key'], keyword) or self._match_keyword(str(c['value']), keyword)]
            if category:
                matched = [c for c in matched if self._match_keyword(c['category'], category)]

            total = len(matched)
            page = matched[offset:offset + limit]

            summary_items = []
            for c in page:
                summary_items.append({
                    "key": c['key'],
                    "value": str(c['value']),
                    "category": c['category'],
                    "type": "config",
                })

            return self._build_result(summary_items, total, keyword or "", offset, limit)

        except Exception as e:
            return ToolResult.err(f"配置搜索失败：{e}")

    def _guess_category(self, key: str) -> str:
        """从键名推测类别"""
        key_lower = key.lower()
        if 'theme' in key_lower or 'color' in key_lower:
            return 'theme'
        if 'model' in key_lower or 'llm' in key_lower or 'api' in key_lower:
            return 'model'
        if 'tool' in key_lower:
            return 'tools'
        if 'session' in key_lower:
            return 'session'
        if 'skill' in key_lower:
            return 'skills'
        return 'general'

    def _format_item(self, item: dict) -> str:
        return f"  ⚙️ {item['key']}  [{item['category']}]  = {item['value']}"


# ---------- SearchSkillsTool ----------

class SearchSkillsTool(SearchTool):
    """技能搜索工具：替代 skill_manager list 的全量返回，支持按名称/描述关键词搜索"""
    name = "search_skills"
    category = "skill_search"
    keywords = ["search", "skills", "技能搜索", "查找技能"]
    model_hint = "当需要搜索已注册的技能（按名称或描述关键词）而非列出全部技能时使用。"
    description = "搜索系统中的技能，支持按名称、描述关键词过滤，返回摘要列表。"
    parameters = {
        "type": "object",
        "properties": {
            "keyword": {"type": "string", "description": "技能名称或描述的关键词", "example": "analyze"},
            "status": {"type": "string", "description": "按启用状态过滤（可选）", "enum": ["enabled", "disabled", "all"], "example": "enabled"},
            "skill_type": {"type": "string", "description": "按技能类型过滤（可选）", "example": "builtin"},
            "limit": {"type": "integer", "description": "每页返回数量，默认 10", "default": 10},
            "offset": {"type": "integer", "description": "分页偏移量，默认 0", "default": 0},
        },
        "required": [],
    }
    item_type_label = "技能"

    def execute(self, keyword: Optional[str] = None, status: Optional[str] = None,
                skill_type: Optional[str] = None, limit: int = 10, offset: int = 0) -> str:
        result = self.execute_structured(keyword, status, skill_type, limit, offset)
        return result.to_json()

    def execute_structured(self, keyword: Optional[str] = None, status: Optional[str] = None,
                           skill_type: Optional[str] = None, limit: int = 10, offset: int = 0) -> ToolResult:
        try:
            from .config import get_config_manager
            config_mgr = get_config_manager()

            all_skills = list(config_mgr.skill_configs.values())

            # 状态过滤
            if status == "enabled":
                all_skills = [s for s in all_skills if s.enabled]
            elif status == "disabled":
                all_skills = [s for s in all_skills if not s.enabled]

            # 类型过滤
            if skill_type:
                all_skills = [s for s in all_skills if self._match_keyword(s.type, skill_type)]

            # 关键词过滤
            if keyword:
                all_skills = [
                    s for s in all_skills
                    if self._match_keyword(s.name, keyword) or self._match_keyword(s.description, keyword)
                ]

            total = len(all_skills)
            page = all_skills[offset:offset + limit]

            summary_items = []
            for s in page:
                icon = "✅" if s.enabled else "❌"
                summary_items.append({
                    "name": s.name,
                    "description": s.description,
                    "type": s.type,
                    "status": "enabled" if s.enabled else "disabled",
                    "label": f"{icon} {s.name}",
                })

            return self._build_result(summary_items, total, keyword or "", offset, limit)

        except Exception as e:
            return ToolResult.err(f"技能搜索失败：{e}")

    def _format_item(self, item: dict) -> str:
        return f"  {item['label']}  [{item['type']}]  {item['description']}"


# ---------- SearchDocumentsTool ----------

class SearchDocumentsTool(SearchTool):
    """文档搜索工具：替代 rag list_documents 的全量返回，支持按文件名关键词搜索"""
    name = "search_documents"
    category = "rag_search"
    keywords = ["search", "documents", "文档搜索", "知识库搜索", "查找文档"]
    model_hint = "当需要搜索向量知识库中的文档（按文件名关键词）而非列出全部文档时使用。"
    description = "搜索 RAG 向量知识库中的文档，支持按文件名关键词过滤，返回摘要列表。"
    parameters = {
        "type": "object",
        "properties": {
            "keyword": {"type": "string", "description": "文件名关键词", "example": "readme"},
            "limit": {"type": "integer", "description": "每页返回数量，默认 10", "default": 10},
            "offset": {"type": "integer", "description": "分页偏移量，默认 0", "default": 0},
        },
        "required": [],
    }
    item_type_label = "文档"

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()
        self._rag_engine = None

    def _get_rag_engine(self):
        if self._rag_engine is None:
            try:
                from .rag.rag_engine import RAGEngine
                persist_dir = self.base_dir / "rag_data"
                persist_dir.mkdir(exist_ok=True)
                self._rag_engine = RAGEngine(persist_directory=persist_dir)
            except Exception as e:
                logger.warning(f"RAG引擎初始化失败: {e}")
                return None
        return self._rag_engine

    def execute(self, keyword: Optional[str] = None, limit: int = 10, offset: int = 0) -> str:
        result = self.execute_structured(keyword, limit, offset)
        return result.to_json()

    def execute_structured(self, keyword: Optional[str] = None, limit: int = 10, offset: int = 0) -> ToolResult:
        try:
            rag_engine = self._get_rag_engine()
            if rag_engine is None:
                return ToolResult.err("RAG引擎初始化失败")

            docs = rag_engine.get_all_documents()

            # 关键词过滤
            if keyword:
                docs = [d for d in docs if self._match_keyword(d.source, keyword)]

            total = len(docs)
            page = docs[offset:offset + limit]

            summary_items = []
            for d in page:
                preview = d.page_content[:80].replace('\n', ' ')
                summary_items.append({
                    "source": d.source,
                    "id": d.id,
                    "preview": preview,
                })

            return self._build_result(summary_items, total, keyword or "", offset, limit)

        except Exception as e:
            return ToolResult.err(f"文档搜索失败：{e}")

    def _format_item(self, item: dict) -> str:
        return f"  📄 {item['source']}\n     预览: {item['preview']}..."


# ---------- SearchMCPToolsTool ----------

class SearchMCPToolsTool(SearchTool):
    """MCP工具搜索工具：替代 mcp_client list_tools 的全量返回，支持按名称/描述关键词搜索"""
    name = "search_mcp_tools"
    category = "mcp_search"
    keywords = ["search", "mcp", "tools", "MCP工具搜索", "查找外部工具"]
    model_hint = "当需要搜索已连接的外部 MCP 服务工具（按名称或描述关键词）而非列出全部工具时使用。"
    description = "搜索已连接的外部 MCP 工具，支持按名称、描述关键词过滤，返回摘要列表。"
    parameters = {
        "type": "object",
        "properties": {
            "keyword": {"type": "string", "description": "工具名称或描述的关键词", "example": "weather"},
            "server": {"type": "string", "description": "按 MCP 服务器名称过滤", "example": "weather_server"},
            "limit": {"type": "integer", "description": "每页返回数量，默认 10", "default": 10},
            "offset": {"type": "integer", "description": "分页偏移量，默认 0", "default": 0},
        },
        "required": [],
    }
    item_type_label = "MCP工具"

    def execute(self, keyword: Optional[str] = None, server: Optional[str] = None,
                limit: int = 10, offset: int = 0) -> str:
        result = self.execute_structured(keyword, server, limit, offset)
        return result.to_json()

    def execute_structured(self, keyword: Optional[str] = None, server: Optional[str] = None,
                           limit: int = 10, offset: int = 0) -> ToolResult:
        try:
            from .mcp_client.client import MCPClientManager
            client_mgr = MCPClientManager()

            tools = client_mgr.list_tools()

            # 服务器过滤
            if server:
                tools = [t for t in tools if self._match_keyword(t.get("server_name", ""), server)]

            # 关键词过滤
            if keyword:
                tools = [
                    t for t in tools
                    if self._match_keyword(t.get("name", ""), keyword)
                    or self._match_keyword(t.get("description", ""), keyword)
                ]

            total = len(tools)
            page = tools[offset:offset + limit]

            summary_items = []
            for t in page:
                summary_items.append({
                    "name": t.get("name", "未知"),
                    "description": t.get("description", "无描述"),
                    "server": t.get("server_name", ""),
                })

            return self._build_result(summary_items, total, keyword or "", offset, limit)

        except Exception as e:
            return ToolResult.err(f"MCP工具搜索失败：{e}")

    def _format_item(self, item: dict) -> str:
        server_info = f" [{item['server']}]" if item.get('server') else ""
        return f"  🔧 {item['name']}{server_info}: {item['description']}"


class OCRTool(Tool):
    """Umi-OCR 本地 HTTP API 集成，支持图片文字识别、二维码识别和混合模式"""
    name = "ocr_recognize"
    category = "vision"
    keywords = ["ocr", "文字识别", "图片识别", "二维码", "QR", "扫描", "识别", "umi", "公式", "latex", "混合", "pdf", "文档"]
    risk_level = "low"
    model_hint = (
        "[何时使用] 需要从图片中提取文字、识别截图内容、读取二维码时使用。"
        "需要 Umi-OCR 本地服务运行中（默认 http://127.0.0.1:1224）。\n"
        "[参数说明]\n"
        "- image_path: 必填，图片文件路径（支持 png/jpg/bmp/webp 等常见格式）\n"
        "- mode: 识别模式，ocr=纯文字识别（默认），qrcode=二维码识别，"
        "mixed=文字+公式混合识别，pdf=PDF文字识别，pdf_mixed=PDF文字+公式混合识别\n"
        "- language: OCR 语言，如 简体中文/繁体中文/English/日本語（默认 简体中文）\n"
        "- data_format: 返回格式，text=纯文本（默认），dict=含位置和置信度\n"
        "- parser: 排版解析，multi_para=多栏按自然段（默认），single_line=单栏逐行，single_code=保留缩进\n"
        "- formula_score_threshold: mixed模式下，置信度低于此阈值的文本块视为公式（默认0.85）\n"
        "- server_url: Umi-OCR 服务地址，默认 http://127.0.0.1:1224\n"
        "[mixed模式说明] 先用Umi-OCR识别全页文字和布局，低置信度区域（通常是公式/符号）"
        "自动裁剪后交给LaTeX-OCR识别为LaTeX代码，最终合并输出。适合含数学公式的学术文档。"
    )
    description = "调用本地 Umi-OCR 服务识别图片中的文字或二维码，支持混合模式自动分流公式区域。"
    parameters = {
        "type": "object",
        "properties": {
            "image_path": {
                "type": "string",
                "description": "图片文件路径（相对于工作目录或绝对路径）",
                "example": "screenshot.png",
            },
            "mode": {
                "type": "string",
                "description": "识别模式：ocr=文字识别，qrcode=二维码，mixed=文字+公式混合，pdf=PDF文字识别，pdf_mixed=PDF+公式",
                "enum": ["ocr", "qrcode", "mixed", "pdf", "pdf_mixed"],
                "default": "ocr",
            },
            "language": {
                "type": "string",
                "description": "OCR 识别语言（仅 ocr/mixed 模式有效）",
                "default": "简体中文",
                "example": "English",
            },
            "data_format": {
                "type": "string",
                "description": "返回数据格式（仅 ocr 模式有效）：text=纯文本，dict=含位置信息",
                "enum": ["text", "dict"],
                "default": "text",
            },
            "parser": {
                "type": "string",
                "description": "排版解析方案（仅 ocr/mixed 模式有效）",
                "enum": ["multi_para", "multi_line", "multi_none",
                         "single_para", "single_line", "single_none", "single_code", "none"],
                "default": "multi_para",
            },
            "formula_score_threshold": {
                "type": "number",
                "description": "mixed模式下公式判定阈值：置信度低于此值的文本块视为公式（0~1，默认0.85）",
                "default": 0.85,
            },
            "server_url": {
                "type": "string",
                "description": "Umi-OCR HTTP 服务地址",
                "default": "http://127.0.0.1:1224",
            },
        },
        "required": ["image_path"],
    }

    # 语言名称 → PaddleOCR/RapidOCR 引擎参数映射
    _LANGUAGE_MAP = {
        "简体中文": "models/config_chinese.txt",
        "繁體中文": "models/config_chinese_cht.txt",
        "繁体中文": "models/config_chinese_cht.txt",
        "english": "models/config_en.txt",
        "English": "models/config_en.txt",
        "日本語": "models/config_japan.txt",
        "日本语": "models/config_japan.txt",
        "한국어": "models/config_korean.txt",
        "韩语": "models/config_korean.txt",
        "Русский": "models/config_cyrillic.txt",
        "俄语": "models/config_cyrillic.txt",
    }

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()
        self._cached_lang_map = None

    def _resolve_language(self, language: str, server_url: str) -> str:
        """将友好语言名映射为引擎参数，若已是引擎格式则直接返回"""
        if language.startswith("models/"):
            return language
        mapped = self._LANGUAGE_MAP.get(language)
        if mapped:
            return mapped
        # 尝试从服务端 get_options 动态获取映射
        try:
            import requests as req_mod
            if self._cached_lang_map is None:
                resp = req_mod.get(f"{server_url.rstrip('/')}/api/ocr/get_options", timeout=5)
                opts = resp.json()
                lang_opt = opts.get("ocr.language", {})
                self._cached_lang_map = {}
                for key, label in lang_opt.get("optionsList", []):
                    self._cached_lang_map[label] = key
                    self._cached_lang_map[key] = key
            return self._cached_lang_map.get(language, language)
        except Exception:
            return language

    def _call_umi_ocr(self, img_b64: str, server_url: str, language: str,
                       data_format: str, parser: str) -> dict:
        """调用 Umi-OCR API，返回原始响应 dict"""
        import requests as req_mod
        url = f"{server_url.rstrip('/')}/api/ocr"
        options = {
            "data.format": data_format,
            "tbpu.parser": parser,
        }
        if language:
            options["ocr.language"] = self._resolve_language(language, server_url)
        payload = {"base64": img_b64, "options": options}
        resp = req_mod.post(url, json=payload, timeout=60)
        resp.raise_for_status()
        return resp.json()

    def _crop_and_latex(self, img_path: Path, boxes: list) -> list:
        """裁剪公式区域并用 LatexOCR 识别，返回 [(box_index, latex_code), ...]"""
        try:
            from PIL import Image
        except ImportError:
            return []
        try:
            from pix2tex.cli import LatexOCR as Pix2TexModel
        except ImportError:
            logger.warning("pix2tex 未安装，mixed 模式无法识别公式")
            return []

        # 懒加载 LatexOCR 模型
        if not hasattr(self, '_latex_model') or self._latex_model is None:
            try:
                self._latex_model = Pix2TexModel()
            except Exception as e:
                logger.warning(f"LatexOCR 模型加载失败: {e}")
                self._latex_model = None
                return []

        try:
            full_img = Image.open(str(img_path))
        except Exception:
            return []

        results = []
        for i, box in enumerate(boxes):
            try:
                # box 格式: [[x1,y1],[x2,y2],[x3,y3],[x4,y4]]
                xs = [p[0] for p in box]
                ys = [p[1] for p in box]
                left = max(0, min(xs) - 5)
                upper = max(0, min(ys) - 5)
                right = min(full_img.width, max(xs) + 5)
                lower = min(full_img.height, max(ys) + 5)
                cropped = full_img.crop((left, upper, right, lower))
                latex = self._latex_model(cropped)
                if latex and latex.strip():
                    results.append((i, latex.strip()))
            except Exception as e:
                logger.debug(f"公式区域 {i} 识别失败: {e}")
                continue
        return results

    def _process_pdf(self, full_path: Path, mode: str, language: str,
                     parser: str, formula_score_threshold: float,
                     server_url: str) -> ToolResult:
        """处理 PDF 文件，逐页转为图片后 OCR"""
        import base64 as b64mod
        import requests as req_mod

        is_mixed = mode in ("pdf_mixed", "mixed")

        try:
            import fitz  # PyMuPDF
        except ImportError:
            return ToolResult.err(
                "缺少 PyMuPDF 库，无法处理 PDF。"
                "请安装：pip install pymupdf"
            )

        try:
            doc = fitz.open(str(full_path))
        except Exception as e:
            return ToolResult.err(f"无法打开 PDF 文件: {e}")

        if doc.page_count == 0:
            doc.close()
            return ToolResult.err("PDF 文件为空")

        all_texts = []
        all_formula_count = 0
        all_text_count = 0
        total_elapsed = 0.0

        for page_num in range(doc.page_count):
            page = doc[page_num]
            # 渲染为图片（150 DPI，适合 OCR）
            mat = fitz.Matrix(150/72, 150/72)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            img_b64 = b64mod.b64encode(img_bytes).decode("ascii")

            actual_format = "dict" if is_mixed else "text"
            try:
                result_data = self._call_umi_ocr(img_b64, server_url, language,
                                                  actual_format, parser)
            except req_mod.ConnectionError:
                doc.close()
                return ToolResult.err(f"无法连接 Umi-OCR 服务 ({server_url})")
            except req_mod.Timeout:
                doc.close()
                return ToolResult.err("Umi-OCR 请求超时（60秒）")
            except Exception as e:
                doc.close()
                return ToolResult.err(f"请求 Umi-OCR 失败: {e}")

            code = result_data.get("code", -1)
            elapsed = result_data.get("time", 0)
            total_elapsed += elapsed

            if code == 101:
                all_texts.append(f"[第 {page_num+1} 页：无文字内容]")
                continue
            if code != 100:
                all_texts.append(f"[第 {page_num+1} 页：识别失败 code={code}]")
                continue

            data = result_data.get("data")

            if is_mixed and isinstance(data, list):
                # 公式分流
                formula_blocks = []
                text_parts = []
                for i, item in enumerate(data):
                    score = item.get("score", 1.0)
                    if score < formula_score_threshold:
                        formula_blocks.append((i, item))
                    else:
                        text_parts.append(f"{item.get('text', '')}{item.get('end', '')}")

                latex_results = {}
                if formula_blocks:
                    boxes = [item["box"] for _, item in formula_blocks if "box" in item]
                    if boxes:
                        # 用 PIL Image 做裁剪（PyMuPDF pixmap 转 PIL）
                        from PIL import Image
                        import io
                        pil_img = Image.open(io.BytesIO(img_bytes))
                        temp_path = full_path.parent / f"__temp_page_{page_num}.png"
                        pil_img.save(str(temp_path))
                        latex_pairs = self._crop_and_latex(temp_path, boxes)
                        for idx, latex in latex_pairs:
                            latex_results[formula_blocks[idx][0]] = latex.strip()
                        temp_path.unlink(missing_ok=True)

                merged = []
                for i, item in enumerate(data):
                    if i in latex_results:
                        merged.append(f"${latex_results[i]}$")
                        all_formula_count += 1
                    else:
                        merged.append(f"{item.get('text', '')}{item.get('end', '')}")
                        all_text_count += 1
                page_text = "".join(merged)
                all_text_count += len(data) - len(formula_blocks)
            else:
                # 普通模式
                if isinstance(data, list):
                    page_text = "".join(f"{item.get('text','')}{item.get('end','')}" for item in data)
                    all_text_count += len(data)
                else:
                    page_text = str(data) if data else ""

            all_texts.append(f"[第 {page_num+1} 页]\n{page_text}")

        page_count = doc.page_count
        doc.close()

        final_text = "\n\n".join(all_texts)
        return ToolResult.ok(
            data={
                "text": final_text,
                "page_count": page_count,
                "formula_count": all_formula_count,
                "text_count": all_text_count,
                "total_elapsed": total_elapsed,
                "image": str(full_path),
                "mode": mode,
            },
            message=f"✅ PDF 识别完成（共 {len(all_texts)} 页，{all_text_count} 文本块，"
                    f"{all_formula_count} 公式块，耗时 {total_elapsed:.1f}s）\n\n{final_text}"
        )

    def execute(self, image_path: str, mode: str = "ocr", language: str = "简体中文",
                data_format: str = "text", parser: str = "multi_para",
                formula_score_threshold: float = 0.85,
                server_url: str = "http://127.0.0.1:1224", **kwargs) -> str:
        result = self.execute_structured(image_path, mode, language, data_format, parser,
                                         formula_score_threshold, server_url)
        return result.to_json()

    def execute_structured(self, image_path: str, mode: str = "ocr", language: str = "简体中文",
                           data_format: str = "text", parser: str = "multi_para",
                           formula_score_threshold: float = 0.85,
                           server_url: str = "http://127.0.0.1:1224", **kwargs) -> ToolResult:
        try:
            import base64 as b64mod
            import requests as req_mod
        except ImportError as e:
            return ToolResult.err(f"缺少依赖库: {e}（需要 requests 库）")

        # 解析图片路径
        full_path = self.base_dir / Path(image_path)
        if not full_path.exists():
            full_path = Path(image_path)
        if not full_path.exists():
            return ToolResult.err(f"图片文件不存在: {image_path}")

        # ===== PDF 模式自动检测 =====
        if mode in ("pdf", "pdf_mixed") or full_path.suffix.lower() in (".pdf",):
            if mode in ("ocr", "mixed", "qrcode"):
                mode = "pdf"  # 自动升级为 PDF 模式
            return self._process_pdf(full_path, mode, language, parser,
                                     formula_score_threshold, server_url)

        # 读取图片并转为 base64
        try:
            img_bytes = full_path.read_bytes()
            img_b64 = b64mod.b64encode(img_bytes).decode("ascii")
        except Exception as e:
            return ToolResult.err(f"读取图片失败: {e}")

        # ===== 二维码模式 =====
        if mode == "qrcode":
            url = f"{server_url.rstrip('/')}/api/qrcode"
            payload = {"base64": img_b64}
            try:
                resp = req_mod.post(url, json=payload, timeout=60)
                resp.raise_for_status()
                result_data = resp.json()
            except req_mod.ConnectionError:
                return ToolResult.err(f"无法连接 Umi-OCR 服务 ({server_url})")
            except req_mod.Timeout:
                return ToolResult.err("Umi-OCR 请求超时（60秒）")
            except Exception as e:
                return ToolResult.err(f"请求 Umi-OCR 失败: {e}")

            code = result_data.get("code", -1)
            if code == 101:
                return ToolResult.ok(data={"raw": result_data}, message="图片中未检测到二维码")
            if code != 100:
                error_msg = result_data.get("data", "未知错误")
                if isinstance(error_msg, list):
                    error_msg = str(error_msg)
                return ToolResult.err(f"二维码识别失败 (code={code}): {error_msg}")

            data = result_data.get("data")
            elapsed = result_data.get("time", 0)
            if isinstance(data, list):
                codes = [{"text": item.get("text", ""), "format": item.get("format", "")} for item in data]
                text_parts = [f"[{c['format']}] {c['text']}" for c in codes]
                return ToolResult.ok(
                    data={"codes": codes, "elapsed": elapsed, "image": str(image_path)},
                    message=f"✅ 识别到 {len(codes)} 个二维码（耗时 {elapsed:.1f}s）\n\n" + "\n".join(text_parts)
                )
            return ToolResult.ok(data={"raw": result_data}, message=f"二维码结果: {data}")

        # ===== OCR / Mixed 模式 =====
        # mixed 模式强制用 dict 格式获取位置信息
        actual_format = "dict" if mode == "mixed" else data_format

        try:
            result_data = self._call_umi_ocr(img_b64, server_url, language, actual_format, parser)
        except req_mod.ConnectionError:
            return ToolResult.err(f"无法连接 Umi-OCR 服务 ({server_url})，请确认 Umi-OCR 已启动并开启 HTTP 服务")
        except req_mod.Timeout:
            return ToolResult.err("Umi-OCR 请求超时（60秒）")
        except Exception as e:
            return ToolResult.err(f"请求 Umi-OCR 失败: {e}")

        code = result_data.get("code", -1)
        if code == 101:
            return ToolResult.ok(data={"raw": result_data}, message="图片中未检测到文字")
        if code != 100:
            error_msg = result_data.get("data", "未知错误")
            if isinstance(error_msg, list):
                error_msg = str(error_msg)
            return ToolResult.err(f"OCR 识别失败 (code={code}): {error_msg}")

        data = result_data.get("data")
        elapsed = result_data.get("time", 0)

        # ===== 纯 OCR 模式 =====
        if mode == "ocr":
            if data_format == "text":
                text = data if isinstance(data, str) else str(data)
                return ToolResult.ok(
                    data={"text": text, "elapsed": elapsed, "image": str(image_path)},
                    message=f"✅ OCR 识别完成（耗时 {elapsed:.1f}s）\n\n{text}"
                )
            else:  # dict
                if isinstance(data, list):
                    lines = [f"{item.get('text', '')}{item.get('end', '')}" for item in data]
                    plain_text = "".join(lines).strip()
                    return ToolResult.ok(
                        data={"blocks": data, "elapsed": elapsed, "image": str(image_path)},
                        message=f"✅ OCR 识别完成（{len(data)} 个文本块，耗时 {elapsed:.1f}s）\n\n{plain_text}"
                    )
                return ToolResult.ok(data={"raw": result_data}, message=f"OCR 结果: {data}")

        # ===== Mixed 模式：文字 + 公式 =====
        if not isinstance(data, list):
            # dict 格式未返回列表，降级为纯文本
            text = data if isinstance(data, str) else str(data)
            return ToolResult.ok(
                data={"text": text, "elapsed": elapsed, "image": str(image_path), "mode": "mixed_fallback"},
                message=f"✅ OCR 识别完成（降级为纯文本，耗时 {elapsed:.1f}s）\n\n{text}"
            )

        # 分流：低置信度块 → 公式，高置信度块 → 文字
        text_blocks = []
        formula_blocks = []  # (index, block_dict)
        for i, item in enumerate(data):
            score = item.get("score", 1.0)
            if score < formula_score_threshold:
                formula_blocks.append((i, item))
            else:
                text_blocks.append(item)

        # 对公式区域裁剪并用 LatexOCR 识别
        latex_results = {}
        if formula_blocks:
            boxes = [item["box"] for _, item in formula_blocks if "box" in item]
            if boxes:
                latex_pairs = self._crop_and_latex(full_path, boxes)
                for idx, latex in latex_pairs:
                    latex_results[formula_blocks[idx][0]] = latex

        # 合并输出
        output_parts = []
        merged_blocks = []
        formula_count = 0
        for i, item in enumerate(data):
            if i in latex_results:
                latex = latex_results[i]
                output_parts.append(f"${latex}$")
                merged_blocks.append({"type": "formula", "latex": latex, "original": item})
                formula_count += 1
            else:
                text = item.get("text", "")
                end = item.get("end", "")
                output_parts.append(f"{text}{end}")
                merged_blocks.append({"type": "text", **item})

        final_text = "".join(output_parts).strip()

        return ToolResult.ok(
            data={
                "text": final_text,
                "blocks": merged_blocks,
                "formula_count": formula_count,
                "text_count": len(data) - formula_count,
                "elapsed": elapsed,
                "image": str(image_path),
                "mode": "mixed",
            },
            message=f"✅ 混合识别完成（{len(data) - formula_count} 文本块 + {formula_count} 公式块，耗时 {elapsed:.1f}s）\n\n{final_text}"
        )


class LatexOCRTool(Tool):
    """LaTeX-OCR (pix2tex) 集成，将数学公式图片转换为 LaTeX 代码"""
    name = "latex_ocr"
    category = "vision"
    keywords = ["latex", "公式", "formula", "数学", "math", "pix2tex", "equation"]
    risk_level = "low"
    model_hint = (
        "[何时使用] 需要将数学公式图片转换为 LaTeX 代码时使用。"
        "支持手写或印刷的数学公式、方程式、符号表达式等。"
        "不支持普通文字识别（用 ocr_recognize）。\n"
        "[参数说明]\n"
        "- image_path: 必填，公式图片文件路径\n"
        "- temperature: 生成温度，越低越确定（默认 0.25）\n"
        "- no_resize: 是否禁用图像自动缩放（默认 False，建议保持）"
    )
    description = "将数学公式图片识别为 LaTeX 代码（基于 pix2tex 深度学习模型）。"
    parameters = {
        "type": "object",
        "properties": {
            "image_path": {
                "type": "string",
                "description": "公式图片文件路径（相对于工作目录或绝对路径）",
                "example": "formula.png",
            },
            "temperature": {
                "type": "number",
                "description": "生成温度，越低结果越确定，越高越多样（0.01~1.0）",
                "default": 0.25,
            },
            "no_resize": {
                "type": "boolean",
                "description": "是否禁用图像自动缩放（通常不需要）",
                "default": False,
            },
        },
        "required": ["image_path"],
    }

    # 类级别模型缓存，避免重复加载
    _model = None
    _model_loading = False

    def __init__(self, base_dir: Optional[Path] = None):
        self.base_dir = base_dir or Path.cwd()

    def _get_model(self):
        """懒加载 LatexOCR 模型（单例）"""
        if LatexOCRTool._model is not None:
            return LatexOCRTool._model
        if LatexOCRTool._model_loading:
            return None
        LatexOCRTool._model_loading = True
        try:
            from pix2tex.cli import LatexOCR as Pix2TexModel
            LatexOCRTool._model = Pix2TexModel()
            return LatexOCRTool._model
        except Exception as e:
            logger.warning(f"LatexOCR 模型加载失败: {e}")
            return None
        finally:
            LatexOCRTool._model_loading = False

    def execute(self, image_path: str, temperature: float = 0.25, no_resize: bool = False, **kwargs) -> str:
        result = self.execute_structured(image_path, temperature, no_resize)
        return result.to_json()

    def execute_structured(self, image_path: str, temperature: float = 0.25, no_resize: bool = False, **kwargs) -> ToolResult:
        # 解析图片路径
        full_path = self.base_dir / Path(image_path)
        if not full_path.exists():
            full_path = Path(image_path)
        if not full_path.exists():
            return ToolResult.err(f"图片文件不存在: {image_path}")

        # 加载模型
        model = self._get_model()
        if model is None:
            return ToolResult.err(
                "LatexOCR 模型加载失败，请确认 pix2tex 已安装（pip install pix2tex）且模型权重已下载"
            )

        # 读取图片
        try:
            from PIL import Image
            img = Image.open(str(full_path))
        except ImportError:
            return ToolResult.err("缺少依赖库 Pillow（pip install Pillow）")
        except Exception as e:
            return ToolResult.err(f"读取图片失败: {e}")

        # 设置温度
        original_temp = model.args.get('temperature', 0.25)
        model.args.temperature = temperature

        # 执行识别
        try:
            latex_code = model(img, resize=not no_resize)
        except Exception as e:
            model.args.temperature = original_temp
            return ToolResult.err(f"LaTeX 识别失败: {e}")
        finally:
            model.args.temperature = original_temp

        if not latex_code or not latex_code.strip():
            return ToolResult.ok(
                data={"latex": "", "image": str(image_path)},
                message="图片中未检测到数学公式"
            )

        return ToolResult.ok(
            data={"latex": latex_code, "image": str(image_path)},
            message=f"✅ LaTeX 识别完成\n\n```latex\n{latex_code}\n```"
        )


class ActivateToolGroupTool(Tool):
    """主动激活工具组：让 LLM 在需要时自行激活专业工具组"""
    name = "activate_tool_group"
    category = "system"
    keywords = ["activate", "tool", "group", "激活", "工具组"]
    model_hint = "当你需要使用当前未激活的专业工具组（如数学计算、动画渲染、定理证明等）时，先调用此工具激活对应组。"
    description = "激活指定工具组，使其工具在后续对话中可用。当需要使用专业工具但当前不可用时调用。"
    parameters = {
        "type": "object",
        "properties": {
            "group": {
                "type": "string",
                "description": "要激活的工具组名称",
                "enum": ["datahub", "scholar", "wolfram", "lean4", "manim",
                         "mathlens", "autoresearch", "feishu", "gt"],
            },
        },
        "required": ["group"],
    }

    def __init__(self, base_dir=None):
        self._base_dir = base_dir
        self._group_mgr = None

    def set_group_manager(self, mgr):
        """注入 ToolGroupManager 实例"""
        self._group_mgr = mgr

    def execute(self, group: str = "") -> str:
        result = self.execute_structured(group)
        return result.to_json()

    def execute_structured(self, group: str = "") -> ToolResult:
        if not group:
            return ToolResult.err("必须指定要激活的工具组名称")

        mgr = self._group_mgr
        if mgr is None:
            return ToolResult.err("工具组管理器未初始化")

        result = mgr.activate_group(group)

        if not result.get("success"):
            return ToolResult.err(
                result.get("error", "激活失败"),
                data={"available_groups": result.get("available_groups", "")}
            )

        tools_list = result.get("tools", [])
        tools_str = ", ".join(tools_list) if tools_list else "（该组暂无可用工具）"

        msg = f"✅ 已激活工具组: {result['label']} ({result['group']})"
        if result.get("was_already_active"):
            msg = f"ℹ️ 工具组 {result['label']} 已处于激活状态"
        msg += f"\n可用工具 ({result['tools_count']}): {tools_str}"
        msg += "\n这些工具将在下一轮对话中可用。"

        return ToolResult.ok(
            data=result,
            message=msg
        )


# === 新增工具：apply_patch, list_code_definition_names, access_mcp_resource, ask_followup_question, browser_action ===


class ApplyPatchTool(Tool):
    name = "apply_patch"
    category = "file_io"
    keywords = ["patch", "diff", "补丁"]
    model_hint = ("[何时使用] 批量修改文件时使用。\n[参数说明]\n- input: V4A 格式 diff 补丁")
    description = "应用 V4A 格式 diff 补丁到文件（支持添加/更新/删除）。"
    parameters = {
        "type": "object",
        "properties": {
            "input": {"type": "string", "description": "【必填】V4A 格式 diff 补丁内容，包含 *** Begin Patch / *** End Patch 标记和 Add/Update/Delete 指令块"},
        },
        "required": ["input"],
    }

    def __init__(self, base_dir=None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, input="") -> str:
        result = self.execute_structured(input)
        return result.to_json()

    def execute_structured(self, input="") -> ToolResult:
        if not input or "*** Begin Patch" not in input:
            return ToolResult.err("补丁格式无效：缺少 *** Begin Patch")
        body = input.split("*** Begin Patch", 1)[1].split("*** End Patch", 1)[0].strip()
        if not body:
            return ToolResult.err("补丁内容为空")
        blocks = re.split(r"\*\*\* (Add|Update|Delete) File: ", body)
        results, errors = [], []
        i = 1
        while i < len(blocks):
            action = blocks[i].strip()
            i += 1
            if i >= len(blocks):
                break
            parts = blocks[i].strip().split("\n", 1)
            fpath_str = parts[0].strip()
            content_text = parts[1].strip() if len(parts) > 1 else ""
            i += 1
            fp = self.base_dir / Path(fpath_str)
            try:
                if action == "Add":
                    if fp.exists():
                        errors.append(f"已存在: {fpath_str}")
                        continue
                    clean = "\n".join(l[2:] if l.startswith("+") else l for l in content_text.split("\n"))
                    fp.parent.mkdir(parents=True, exist_ok=True)
                    fp.write_text(clean, encoding="utf-8")
                    results.append(f"+ {fpath_str}")
                elif action == "Delete":
                    if not fp.exists():
                        errors.append(f"不存在: {fpath_str}")
                        continue
                    fp.unlink()
                    results.append(f"- {fpath_str}")
                elif action == "Update":
                    if not fp.exists():
                        errors.append(f"不存在: {fpath_str}")
                        continue
                    orig = fp.read_text(encoding="utf-8")
                    chunks = self._parse_chunks(content_text.split("\n"))
                    current = orig
                    applied = 0
                    for ctx, olds, news in chunks:
                        old_t = "\n".join(olds)
                        if old_t and old_t in current:
                            new_t = "\n".join(news) if news else ""
                            current = current.replace(old_t, new_t, 1)
                            applied += 1
                    if applied > 0:
                        fp.write_text(current, encoding="utf-8")
                        results.append(f"~ {fpath_str} ({applied})")
                    else:
                        errors.append(f"无法匹配: {fpath_str}")
            except Exception as e:
                errors.append(f"{fpath_str}: {e}")
        if not results and errors:
            return ToolResult.err("\n".join(errors))
        msg = "\n".join(results)
        if errors:
            msg += "\n⚠ " + "\n".join(errors[:3])
        return ToolResult.ok(data={"results": results, "errors": errors}, message=msg)

    def _parse_chunks(self, lines):
        chunks, i = [], 0
        while i < len(lines):
            if lines[i].startswith("@@"):
                ctx = lines[i][2:].strip()
                i += 1
                olds, news = [], []
                while i < len(lines) and (lines[i].startswith("-") or lines[i].startswith("+") or not lines[i].strip()):
                    if lines[i].startswith("-"):
                        olds.append(lines[i][2:])
                    elif lines[i].startswith("+"):
                        news.append(lines[i][2:])
                    i += 1
                if olds or news:
                    chunks.append((ctx, olds, news))
            else:
                i += 1
        return chunks


class ListCodeDefinitionNamesTool(Tool):
    name = "list_code_definition_names"
    category = "code_search"
    keywords = ["code", "definitions", "结构"]
    model_hint = "[何时使用] 了解代码库结构时使用。\n[参数说明]\n- path: 目录\n- depth: 深度"
    description = "列出目录中源代码文件的定义名称（类、函数、方法等）。"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "【必填】要扫描的目录路径（相对于工作目录）"},
            "depth": {"type": "integer", "description": "递归扫描深度（默认 3，设为 0 仅扫描当前目录）"},
        },
        "required": ["path"],
    }

    def __init__(self, base_dir=None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, path="", depth=3) -> str:
        return self.execute_structured(path, depth).to_json()

    def execute_structured(self, path="", depth=3) -> ToolResult:
        if not path:
            return ToolResult.err("必须指定目录路径")
        sd = self.base_dir / Path(path)
        if not sd.exists() or not sd.is_dir():
            return ToolResult.err(f"目录不存在: {path}")
        patterns = {
            ".py": r'^(?:async\s+)?(?:def|class)\s+(\w+)',
            ".js": r'(?:export\s+)?(?:class|function|const|let|var)\s+(\w+)',
            ".ts": r'(?:export\s+)?(?:class|interface|function|const|let|var|type|enum)\s+(\w+)',
            ".tsx": r'(?:export\s+)?(?:class|interface|function|const)\s+(\w+)',
            ".jsx": r'(?:export\s+)?(?:class|function|const)\s+(\w+)',
            ".java": r'(?:class|interface|enum)\s+(\w+)',
            ".go": r'^func\s+(\w+)',
            ".rs": r'^(?:pub\s+)?(?:fn|struct|enum|trait|impl|mod)\s+(\w+)',
        }
        excluded = {".git", "__pycache__", "node_modules", ".venv", "venv", "dist", "build", "target"}
        results, fc = {}, 0
        for root, dirs, files in os.walk(str(sd)):
            dirs[:] = [d for d in dirs if d not in excluded and not d.startswith(".")]
            if len(Path(root).relative_to(sd).parts) >= depth:
                dirs.clear()
            for f in sorted(files):
                ext = Path(f).suffix.lower()
                if ext not in patterns:
                    continue
                try:
                    ct = (Path(root) / f).read_text(encoding="utf-8", errors="replace")
                    defs = set()
                    for m in re.finditer(patterns[ext], ct, re.MULTILINE):
                        defs.add(m.group(1))
                    if defs:
                        rp = str((Path(root) / f).relative_to(self.base_dir))
                        results[rp] = sorted(defs)
                        fc += 1
                except Exception:
                    continue
        if not results:
            return ToolResult.ok(data={"definitions": {}}, message="未找到代码定义")
        total = sum(len(v) for v in results.values())
        lines = [f"发现 {total} 个定义，{fc} 个文件:"]
        for fp, defs in sorted(results.items()):
            lines.append(f"  {fp}: {', '.join(defs[:10])}")
        return ToolResult.ok(
            data={"definitions": results, "total_files": fc, "total_definitions": total},
            message="\n".join(lines)
        )


class AccessMCPResourceTool(Tool):
    name = "access_mcp_resource"
    category = "mcp"
    keywords = ["mcp", "resource", "资源"]
    model_hint = ("[何时使用] 读取 MCP 资源时使用。\n[参数说明]\n- server_name: 服务器名\n- uri: 资源 URI")
    description = "访问已连接 MCP 服务器的只读资源。"
    parameters = {
        "type": "object",
        "properties": {
            "server_name": {"type": "string", "description": "【必填】MCP 服务器名称（通过 mcp_client list 或 search_mcp_tools 查看可用服务器）"},
            "uri": {"type": "string", "description": "【必填】资源 URI，如 file:///path 或 custom://resource/path"},
        },
        "required": ["server_name", "uri"],
    }

    def __init__(self, base_dir=None):
        self.base_dir = base_dir or Path.cwd()

    def execute(self, server_name="", uri="") -> str:
        return self.execute_structured(server_name, uri).to_json()

    def execute_structured(self, server_name="", uri="") -> ToolResult:
        if not server_name or not uri:
            return ToolResult.err("必须指定 server_name 和 uri")
        try:
            from ..mcp_server import get_client_manager
            mgr = get_client_manager()
            if not mgr:
                return ToolResult.err("MCP 客户端不可用")
            sessions = getattr(mgr, 'sessions', {}) or getattr(mgr, '_sessions', {})
            s = sessions.get(server_name)
            if not s:
                return ToolResult.err(f"未找到服务器: {server_name}")
            client = getattr(s, 'client', None) or getattr(s, '_client', None)
            if not client:
                return ToolResult.err(f"服务器 {server_name} 无客户端")
            if hasattr(client, 'read_resource'):
                r = client.read_resource(uri)
            elif hasattr(client, 'resources'):
                r = client.resources.read(uri)
            else:
                return ToolResult.err("不支持资源读取")
            return ToolResult.ok(data={"server": server_name, "uri": uri}, message=str(r)[:3000])
        except Exception as e:
            return ToolResult.err(f"访问资源失败: {e}")


class AskFollowupQuestionTool(Tool):
    name = "ask_followup_question"
    category = "interaction"
    keywords = ["ask", "question", "追问"]
    model_hint = ("[何时使用] 需要澄清问题时使用。\n[参数说明]\n- question: 问题\n- options: 选项列表")
    description = "向用户提问以获取额外信息。支持预设选项。"
    parameters = {
        "type": "object",
        "properties": {
            "question": {"type": "string", "description": "【必填】要向用户提问的问题文本"},
            "options": {"type": "array", "items": {"type": "string"}, "description": "预置选项列表（最多5个），用户可直接选择作答"},
        },
        "required": ["question"],
    }

    def execute(self, question="", options=None) -> str:
        return self.execute_structured(question, options).to_json()

    def execute_structured(self, question="", options=None) -> ToolResult:
        if not question:
            return ToolResult.err("必须提供问题")
        msg = f"❓ {question}"
        if options:
            for i, o in enumerate(options[:5], 1):
                msg += f"\n  {i}. {o}"
        return ToolResult.ok(data={"question": question}, message=msg)


class BrowserActionTool(Tool):
    name = "browser_action"
    category = "automation"
    keywords = ["browser", "浏览器", "web"]
    model_hint = ("[何时使用] 需要网页交互时使用。\n[参数说明]\n- action: launch/click/type/scroll/close\n- url/coordinate/text: 对应参数")
    description = "用 Playwright 控制浏览器（启动/点击/输入/滚动/关闭）。"
    parameters = {
        "type": "object",
        "properties": {
            "action": {"type": "string", "enum": ["launch", "click", "type", "scroll_down", "scroll_up", "close"], "description": "【必填】浏览器操作类型：launch=打开URL, click=点击坐标, type=输入文本, scroll_down/scroll_up=滚动, close=关闭浏览器"},
            "url": {"type": "string", "description": "要打开的网页 URL（launch 操作必填），如 https://example.com"},
            "coordinate": {"type": "string", "description": "点击坐标（click 操作必填），格式为 \"x,y\" 如 \"500,300\""},
            "text": {"type": "string", "description": "要输入的文本（type 操作必填），如搜索关键词或表单内容"},
        },
        "required": ["action"],
    }
    _instance = None

    def execute(self, action="", url="", coordinate="", text="") -> str:
        return self.execute_structured(action, url, coordinate, text).to_json()

    def execute_structured(self, action="", url="", coordinate="", text="") -> ToolResult:
        try:
            if action == "launch":
                return self._launch(url)
            elif action == "click":
                return self._click(coordinate)
            elif action == "type":
                return self._type(text)
            elif action in ("scroll_down", "scroll_up"):
                return self._scroll(action == "scroll_down")
            elif action == "close":
                return self._close()
            return ToolResult.err(f"未知操作: {action}")
        except Exception as e:
            return ToolResult.err(f"失败: {e}")

    def _get_page(self):
        if not hasattr(self, '_b') or self._b is None:
            return None
        return self._b

    def _launch(self, url):
        if not url:
            return ToolResult.err("launch 需要 url")
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            return ToolResult.err("需安装 playwright")
        self._close()
        p = sync_playwright().start()
        br = p.chromium.launch(headless=False)
        pg = br.new_page(viewport={"width": 1280, "height": 720})
        pg.goto(url, wait_until="domcontentloaded", timeout=30000)
        BrowserActionTool._instance = (p, br, pg)
        self._b = pg
        return ToolResult.ok(data={"url": url, "title": pg.title()}, message=f"已启动: {url}\n标题: {pg.title()}")
    def _click(self, coord):
        if self._b is None:
            return ToolResult.err("先 launch")
        x, y = int(coord.split(",")[0]), int(coord.split(",")[1])
        self._b.mouse.click(x, y)
        return ToolResult.ok(message=f"点击 ({x},{y})")
    def _type(self, text):
        if self._b is None:
            return ToolResult.err("先 launch")
        self._b.keyboard.type(text)
        return ToolResult.ok(message=f"输入: {text[:30]}")
    def _scroll(self, down):
        if self._b is None:
            return ToolResult.err("先 launch")
        self._b.evaluate(f"window.scrollBy(0, {'window.innerHeight' if down else '-window.innerHeight'})")
        return ToolResult.ok(message=f"滚动{'下' if down else '上'}")
    def _close(self):
        try:
            inst = BrowserActionTool._instance
            if inst:
                p, br, pg = inst
                pg.close()
                br.close()
                p.stop()
                BrowserActionTool._instance = None
            self._b = None
        except Exception:
            pass
        return ToolResult.ok(message="已关闭")


def get_tools(base_dir: Optional[Path] = None, enabled_only: bool = False) -> List[Tool]:
    """获取所有工具
    Args:
        base_dir: 基础目录
        enabled_only: 是否只返回启用的工具
    """
    all_tools = [
        ReadFileTool(base_dir),
        WriteFileTool(base_dir),
        EditFileTool(base_dir),
        ListDirectoryTool(base_dir),
        GrepTool(base_dir),
        GlobTool(base_dir),
        CalculateTool(),
        WebSearchTool(),
        FetchUrlTool(),
        AnalyzePaperTool(base_dir),
        OpenFileTool(base_dir),
        DiffFilesTool(base_dir),
        MoveFileTool(base_dir),
        CopyFileTool(base_dir),
        FileInfoTool(base_dir),
        CLITool(base_dir),
        ConfigTool(base_dir),
        SkillTool(base_dir),
        RAGTool(base_dir),
        SandboxTool(base_dir),
        MCPClientTool(base_dir),
        TerminalTool(base_dir),
        WorkflowTool(base_dir),
        ApprovalTool(base_dir),
        ExtensionSkillTool(base_dir),
        SessionTool(base_dir),
        # SearchTools
        SearchFilesTool(base_dir),
        SearchCoursesTool(),
        SearchSessionsTool(base_dir),
        SearchConfigsTool(base_dir),
        SearchSkillsTool(),
        SearchDocumentsTool(base_dir),
        SearchMCPToolsTool(),
        # 工具组激活
        ActivateToolGroupTool(base_dir),
        # OCR 识别
        OCRTool(base_dir),
        # LaTeX 公式识别
        LatexOCRTool(base_dir),
        # Cline 兼容工具（5个）
        ApplyPatchTool(base_dir),
        ListCodeDefinitionNamesTool(base_dir),
        AccessMCPResourceTool(base_dir),
        AskFollowupQuestionTool(),
        BrowserActionTool(),
    ]

    try:
        from .scholar.server import ScholarMCPServer
        scholar_server = ScholarMCPServer()
        all_tools.extend(scholar_server.get_scholar_tools())
    except Exception:
        pass

    try:
        from .automation.engine import get_automation_engine
        auto_engine = get_automation_engine(base_dir)
        all_tools.extend(auto_engine.get_automation_tools() if hasattr(auto_engine, 'get_automation_tools') else [])
    except Exception:
        pass
    
    if enabled_only:
        try:
            from .config import get_config_manager
            config_mgr = get_config_manager()
            # 确保默认配置存在
            config_mgr.init_default_tool_configs(all_tools)
            enabled_names = [t.name for t in config_mgr.get_enabled_tools()]
            return [tool for tool in all_tools if tool.name in enabled_names]
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"获取启用的工具失败，返回所有工具: {e}")
            return all_tools
    return all_tools


# ──────────────────────────────────────────────
# MCP 远程服务代理工具
# ──────────────────────────────────────────────

class MCPServiceTool(Tool):
    """远程 MCP 服务工具代理 — 每个远程工具自动生成一个本地 Tool 实例

    将远程 MCP 服务的工具包装为本地 Tool，Agent 可直接调用，
    调用被转发到远程 MCP 服务。
    """

    # 这些类属性会被实例属性覆盖
    name = "mcp_service_proxy"
    description = "远程 MCP 服务代理"
    parameters = {"type": "object", "properties": {}}
    category = "mcp_service"
    keywords = ["mcp", "远程", "搜索"]
    risk_level = "low"
    model_hint = ""

    def __init__(
        self,
        tool_name: str,
        tool_description: str,
        tool_parameters: Dict[str, Any],
        server_name: str,
        client_manager: Any = None,
        service_type: str = "mcp",
        rest_url: str = "",
        rest_headers: Dict[str, str] = None,
        rest_method: str = "POST",
        request_transform: str = "",
    ):
        # 动态设置实例属性（覆盖类属性）
        self.name = tool_name
        self.description = tool_description
        self.parameters = tool_parameters
        self.category = "mcp_service"
        self.keywords = ["mcp", server_name]
        self._server_name = server_name
        self._client_manager = client_manager
        self._service_type = service_type  # "mcp" or "rest"
        self._rest_url = rest_url
        self._rest_headers = rest_headers or {}
        self._rest_method = rest_method
        self._request_transform = request_transform

    def execute(self, **kwargs) -> str:
        if self._service_type == "rest":
            return self._execute_rest(**kwargs)
        return self._execute_mcp(**kwargs)

    def _execute_mcp(self, **kwargs) -> str:
        """通过 MCP 协议调用远程工具"""
        try:
            result = self._client_manager.call_tool(self.name, kwargs)
            if isinstance(result, dict):
                # MCP 标准响应格式
                content = result.get("content", [])
                if isinstance(content, list):
                    parts = []
                    for item in content:
                        if isinstance(item, dict):
                            text = item.get("text", "")
                            if text:
                                parts.append(text)
                    if parts:
                        combined = "\n".join(parts)
                        return combined[:MAX_TOOL_RESULT_CHARS]
                # 回退：直接 JSON 序列化
                return json.dumps(result, ensure_ascii=False, indent=2)[:MAX_TOOL_RESULT_CHARS]
            return str(result)[:MAX_TOOL_RESULT_CHARS]
        except Exception as e:
            logger.error(f"MCP 服务工具 '{self.name}' 调用失败: {e}")
            return json.dumps({"success": False, "error": f"MCP 服务调用失败: {e}"}, ensure_ascii=False)

    def _execute_rest(self, **kwargs) -> str:
        """通过 REST API 直接调用远程服务"""
        import urllib.request
        import urllib.error

        if not self._rest_url:
            return json.dumps({"success": False, "error": f"REST 服务 '{self._server_name}' 未配置 URL"}, ensure_ascii=False)

        try:
            # 应用请求转换器
            request_body = self._transform_request(kwargs)

            data = json.dumps(request_body, ensure_ascii=False).encode("utf-8")
            req = urllib.request.Request(
                self._rest_url,
                data=data,
                headers={**self._rest_headers, "Content-Type": "application/json"},
                method=self._rest_method,
            )
            with urllib.request.urlopen(req, timeout=30) as resp:
                result = resp.read().decode("utf-8")
                # 尝试解析为 JSON 美化输出
                try:
                    parsed = json.loads(result)
                    # 百度搜索特殊处理：提取 references
                    if "references" in parsed:
                        refs = parsed["references"]
                        # 移除 snippet 字段减少输出
                        for item in refs:
                            item.pop("snippet", None)
                        return json.dumps(refs, ensure_ascii=False, indent=2)[:MAX_TOOL_RESULT_CHARS]
                    return json.dumps(parsed, ensure_ascii=False, indent=2)[:MAX_TOOL_RESULT_CHARS]
                except (json.JSONDecodeError, TypeError):
                    return result[:MAX_TOOL_RESULT_CHARS]
        except urllib.error.HTTPError as e:
            error_body = ""
            try:
                error_body = e.read().decode("utf-8", errors="replace")
            except Exception:
                pass
            return json.dumps({
                "success": False,
                "error": f"HTTP {e.code}: {e.reason}",
                "detail": error_body[:500],
            }, ensure_ascii=False)
        except urllib.error.URLError as e:
            return json.dumps({"success": False, "error": f"URL Error: {e.reason}"}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"success": False, "error": f"REST 服务调用失败: {e}"}, ensure_ascii=False)

    def _transform_request(self, kwargs: Dict[str, Any]) -> Dict[str, Any]:
        """根据 request_transform 转换请求体"""
        if self._request_transform == "baidu_search":
            return self._transform_baidu_search(kwargs)
        elif self._request_transform == "bocha_search":
            return self._transform_bocha_search(kwargs)
        # 默认：直接传递参数
        return kwargs

    @staticmethod
    def _transform_baidu_search(params: Dict[str, Any]) -> Dict[str, Any]:
        """将简单的 {query, count, freshness} 转换为百度 AI Search API 格式"""
        import re as _re
        from datetime import datetime, timedelta

        query = params.get("query", "")
        count = min(max(int(params.get("count", 10)), 1), 50)
        freshness = params.get("freshness", "")

        search_filter = {}
        if freshness:
            current_time = datetime.now()
            end_date = (current_time + timedelta(days=1)).strftime("%Y-%m-%d")
            if freshness in ("pd", "pw", "pm", "py"):
                delta_map = {"pd": 1, "pw": 6, "pm": 30, "py": 364}
                start_date = (current_time - timedelta(days=delta_map[freshness])).strftime("%Y-%m-%d")
                search_filter = {"range": {"page_time": {"gte": start_date, "lt": end_date}}}
            elif _re.match(r'\d{4}-\d{2}-\d{2}to\d{4}-\d{2}-\d{2}', freshness):
                start_date = freshness.split("to")[0]
                end_date = freshness.split("to")[1]
                search_filter = {"range": {"page_time": {"gte": start_date, "lt": end_date}}}

        return {
            "messages": [{"content": query, "role": "user"}],
            "search_source": "baidu_search_v2",
            "resource_type_filter": [{"type": "web", "top_k": count}],
            "search_filter": search_filter,
        }

    @staticmethod
    def _transform_bocha_search(params: Dict[str, Any]) -> Dict[str, Any]:
        """将简单的 {query, count, freshness} 转换为博查 AI Web Search API 格式"""
        query = params.get("query", "")
        count = min(max(int(params.get("count", 10)), 1), 50)
        freshness = params.get("freshness", "noLimit")
        summary = params.get("summary", True)

        return {
            "query": query,
            "count": count,
            "freshness": freshness,
            "summary": bool(summary),
        }


def _resolve_api_key(env_var: str) -> Optional[str]:
    """从环境变量安全读取 API Key"""
    import os
    return os.environ.get(env_var)


def _load_env_file():
    """加载 .env 文件到环境变量（如果尚未加载）"""
    import os
    env_path = Path(__file__).parent.parent / ".env"
    if not env_path.exists():
        return
    try:
        for line in env_path.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if "=" in line:
                key, _, value = line.partition("=")
                key = key.strip()
                value = value.strip().strip('"').strip("'")
                if key and key not in os.environ:
                    os.environ[key] = value
    except Exception as e:
        logger.debug(f"加载 .env 文件失败: {e}")


def load_mcp_service_tools(config_path: Optional[Path] = None) -> List[Tool]:
    """从 mcp_servers.json 加载远程 MCP 服务，返回代理工具列表

    支持两种服务类型：
    - service_type: "mcp" — 标准 MCP 协议服务（Streamable HTTP/SSE/Stdio）
    - service_type: "rest" — 普通 REST API 服务（直接 HTTP 请求）

    配置文件格式:
    {
      "servers": {
        "baidu-search": {
          "service_type": "rest",
          "url": "https://...",
          "headers": {"Authorization": "Bearer ${BAIDU_API_KEY}"},
          "enabled": true,
          "tools": [{"name": "baidu_search", "description": "...", "inputSchema": {...}}],
          "api_key_env": "BAIDU_API_KEY"
        }
      }
    }
    """
    # 先加载 .env 文件中的环境变量
    _load_env_file()

    if config_path is None:
        config_path = Path(__file__).parent / "mcp_servers.json"

    if not config_path.exists():
        return []

    try:
        config_data = json.loads(config_path.read_text(encoding="utf-8"))
    except Exception as e:
        logger.error(f"读取 MCP 服务配置失败: {e}")
        return []

    servers = config_data.get("servers", {})
    if not servers:
        return []

    # MCP 客户端管理器（仅 MCP 协议服务需要）
    mcp_manager = None
    try:
        from .mcp_client.client import MCPClientManager
        from .mcp_client.transport import TransportConfig
        mcp_manager = MCPClientManager()
    except ImportError:
        logger.warning("MCP 客户端模块不可用，仅支持 REST 类型服务")

    tools: List[Tool] = []

    for server_name, server_config in servers.items():
        if not server_config.get("enabled", True):
            logger.info(f"服务 '{server_name}' 已禁用，跳过")
            continue

        service_type = server_config.get("service_type", "mcp")
        url = server_config.get("url", "")

        # 解析 headers 中的 ${ENV_VAR} 占位符
        raw_headers = server_config.get("headers", {})
        resolved_headers = _resolve_headers(raw_headers, server_name)

        # ── MCP 协议服务：尝试连接发现工具 ──
        if service_type == "mcp" and mcp_manager and url:
            transport_type = server_config.get("transport", "http")
            try:
                transport_config = TransportConfig(
                    type=transport_type,
                    url=url,
                    headers=resolved_headers,
                    timeout=server_config.get("timeout", 30),
                )
                connected = mcp_manager.register(server_name, transport_config)
                if connected:
                    logger.info(f"MCP 服务 '{server_name}' 连接成功，发现 {mcp_manager.get_state(server_name).tool_count} 个工具")
                else:
                    state = mcp_manager.get_state(server_name)
                    logger.warning(f"MCP 服务 '{server_name}' 连接失败: {state.error if state else 'unknown'}")
            except Exception as e:
                logger.warning(f"MCP 服务 '{server_name}' 连接异常: {e}")

        # ── 从配置中的 tools 字段创建代理工具 ──
        declared_tools = server_config.get("tools", [])
        for tool_def in declared_tools:
            tool_name = tool_def.get("name", "")
            if not tool_name:
                continue

            # MCP 协议服务：优先使用远程发现的 schema
            tool_description = tool_def.get("description", "")
            tool_parameters = tool_def.get("inputSchema", {"type": "object", "properties": {}})

            if service_type == "mcp" and mcp_manager:
                for remote_tool in mcp_manager.list_tools():
                    if remote_tool.get("name") == tool_name:
                        tool_description = remote_tool.get("description", tool_description)
                        tool_parameters = remote_tool.get("inputSchema", tool_parameters)
                        break

            proxy = MCPServiceTool(
                tool_name=tool_name,
                tool_description=tool_description,
                tool_parameters=tool_parameters,
                server_name=server_name,
                client_manager=mcp_manager if service_type == "mcp" else None,
                service_type=service_type,
                rest_url=url if service_type == "rest" else "",
                rest_headers=resolved_headers if service_type == "rest" else {},
                rest_method=server_config.get("method", "POST") if service_type == "rest" else "POST",
                request_transform=server_config.get("request_transform", "") if service_type == "rest" else "",
            )
            tools.append(proxy)
            logger.info(f"已注册 {service_type.upper()} 代理工具: {tool_name} (来自 {server_name})")

    return tools


def _resolve_headers(raw_headers: Dict[str, str], server_name: str = "") -> Dict[str, str]:
    """解析 headers 中的 ${ENV_VAR} 占位符"""
    import re as _re
    resolved = {}
    for key, value in raw_headers.items():
        if isinstance(value, str) and "${" in value:
            def _replace_env(match):
                env_var = match.group(1)
                api_key = _resolve_api_key(env_var)
                if api_key:
                    return api_key
                logger.warning(f"服务 '{server_name}': 环境变量 '{env_var}' 未设置，请设置后重试")
                return match.group(0)
            resolved[key] = _re.sub(r'\$\{(\w+)\}', _replace_env, value)
        else:
            resolved[key] = value
    return resolved
